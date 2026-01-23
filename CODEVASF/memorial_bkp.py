#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Gera memorial sintetico (DXF + DOC) a partir de polilinhas/arcos em um DXF.

Para cada polilinha fechada nas layers escolhidas:
- Segmentos sem bulge viram retas (azimute/distância usuais).
- Segmentos com bulge viram arcos (usa a corda para azimute/distância e calcula raio,
  ângulo central e desenvolvimento).

Saídas:
  * DXF com vértices numerados, tabela e ordinates.
  * Documentos DOC (um por polígono) com memorial descritivo simplificado.
  * DXFs individuais por propriedade mostrando vizinhos imediatos.
"""
# Requirements to run this script:
# - Python 3.10+ environment with geopandas, shapely, ezdxf, numpy, pandas, fiona e dependências nativas instaladas.
# - python-docx instalado (obrigatório para gerar os arquivos DOCX).
# - Input DXF available at H:\4 - PROCESSAMENTO\10 MEMORIAL\MEMORIAL.dxf.
# - Support shapefiles in the same folder: PONTO_PROPRIEDADE.shp (atributos) e ESTRADA.shp (geometria da estrada).
# - DOCX template stored at H:\4 - PROCESSAMENTO\10 MEMORIAL\TEMPLATE.docx.
# - DXF template stored at H:\4 - PROCESSAMENTO\10 MEMORIAL\TEMPLATE.dxf.
# - Recommended to run inside the geo_env conda environment configured for GIS workloads.

from __future__ import annotations

import math
import unicodedata
import re
import geopandas as gpd
import rasterio
import geomag
from shapely.geometry import Point, LineString, Polygon
from shapely.geometry.base import BaseGeometry
from datetime import datetime
from math import atan, tan, sin, radians, degrees
from typing import cast

try:
    from docx import Document  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore
except ImportError:  # python-docx not installed
    Document = None  # type: ignore
    OxmlElement = None  # type: ignore
    Paragraph = None  # type: ignore

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import ezdxf

# ---------------------------------------------------------------------------
# CONFIGURAÇÕES GERAIS
# ---------------------------------------------------------------------------

DXF_INPUT_PATH = Path(r"G:\4 - PROCESSAMENTO\10 MEMORIAL\TEMPLATE.dxf")  # DIRETORIO DOS DADOS
DOC_OUTPUT_BASE = DXF_INPUT_PATH.parent

DOC_TEMPLATE_PATH = Path(r"G:\4 - PROCESSAMENTO\10 MEMORIAL\TEMPLATE.docx")
PROPERTIES_SHP_PATH = Path(r"G:\4 - PROCESSAMENTO\10 MEMORIAL\PROPRIEDADES.shp")
DXF_TEMPLATE_PATH = Path(r"G:\4 - PROCESSAMENTO\10 MEMORIAL\TEMPLATE.dxf")
MDT_PATH = Path(r"G:\4 - PROCESSAMENTO\05 MDT\xingo_inteiro.tif")
GEOD_CRS = "EPSG:31984"
WMM_MODEL = "WMM2025"  # alvo 2024-2029; cai para WMM2020 se indisponível
MUNICIPIOS_BA = Path(r"G:\4 - PROCESSAMENTO\10 MEMORIAL\BA_Municipios_2024.shp")
MUNICIPIOS_SE = Path(r"G:\4 - PROCESSAMENTO\10 MEMORIAL\SE_Municipios_2024.shp")

LAYERS_TO_PROCESS: Optional[List[str]] = None

TABLE_PLACEMENT_MODE = "stage"
TABLE_GRID_COLS = 3
TABLE_COLUMN_GAP = 20.0
TABLE_ROW_GAP = 10.0
TABLE_STAGE_XOFFSET = 60.0
TABLE_STAGE_YOFFSET = 20.0
TABLE_SPACING_SCALE = 1.5

VERTEX_NUMBERING_MODE = "per_poly"  # "per_poly" (resetando) ou "global"
PROPERTY_NEIGHBOR_SPAN = 0
SELECT_PROPERTIES_INTERACTIVE = True  # Permite escolher propriedades a exportar (Enter = todas)
VERTEX_MIN_SPACING = 0.01  # 1 cm: distância mínima para considerar novo vértice
GENERATE_DOCS = True  # Controla geração de memorial em DOCX
GENERATE_DXFS = True  # Controla geração de plantas individuais em DXF

# ---------------- Configurações de layers/blocos ----------------

STYLE_TEXTO = "Arial"
BLOCO_VERTICE = "VERTICE"

LAYER_VERTICE_PTO = "VERTICE"
LAYER_VERTICE_TXT = "VERTICE"
LAYER_TABELA = "TABELA"
LAYER_ORDINATE = "TOP_AZIMUTE"
CARIMBO_LAYER = "CAD_CARIMBO"

VERTEX_RADIUS = 2.5
TEMPLATE_TABLE_MARGIN = 5.0
PRIMARY_LAYOUT_NAME = "A2"
SECONDARY_LAYOUT_NAME = "A2"
TABLE_ANCHOR_NO_ARC_X = 29.000   # Âncora padrão da tabela (A2)
TABLE_ANCHOR_NO_ARC_Y = 409.000
TABLE_SECONDARY_ANCHOR_NO_ARC_X = 29.000  # Âncora secundária padrão (A2)
TABLE_SECONDARY_ANCHOR_NO_ARC_Y = 409.000
TABLE_LAYOUT_ANCHORS = {
    "A1": (29.000, 575.500),
    "A2": (29.000, 409.000),
    "A3": (30.628, 282.779),
}
# Mantém tabela inteira na folha 1 (não dividir em secundária)
TABLE_SECONDARY_THRESHOLD = 10_000
DEBUG_MODE = True

# Posicoes do carimbo A2 (camada CAD_CARIMBO)
CARIMBO_DEFAULT_LAYOUT = "A2"
CARIMBO_POSITIONS = {
    "A2": {
        "prop": (447.329, 171.001),
        "owner_main": (448.726, 165.001),
        "owner_aux": (539.256, 73.364),
        "doc_main": (437.370, 159.001),
        "doc_aux": (539.256, 68.552),
        "date": (426.693, 111.001),
        "area": (440.216, 86.931),
        "perimeter": (471.210, 86.931),
        # Folha (sobrepõe valor fixo do template)
        "folha": (569.385, 186.983),
        "escala": (528.731, 117.001),
        "lat": (563.769, 339.359),
        "lon": (563.769, 335.359),
        "convergencia": (515.298, 343.359),
        "declinacao": (515.298, 339.359),
        "variacao": (515.298, 335.359),
        "municipio": (444.708, 153.001),
        "estado": (536.122, 153.001),
        "comarca": (444.708, 147.001),
        "cartorio": (439.246, 141.001),
        "cns": (426.023, 129.001),
    },
    "A3": {
        "prop": (273.329, 143.388),
        "owner_main": (274.727, 137.388),
        "owner_aux": (365.257, 71.886),
        "doc_main": (263.371, 131.388),
        "doc_aux": (365.257, 67.086),
        "date": (343.952, 113.465),
        "area": (270.408, 78.203),
        "perimeter": (300.709, 78.203),
        "folha": (396.345, 155.870),
        "escala": (350.919, 101.388), 
        "lat": (390.742, 239.009),
        "lon": (390.742, 235.009),
        "convergencia": (342.272, 243.009),
        "declinacao": (342.272, 239.009), 
        "variacao": (342.272, 235.009),
        "municipio": (270.708, 125.388),
        "estado": (358.309, 119.388),
        "comarca": (269.977, 119.388),
        "cartorio": (356.505, 143.388),
        "cns": (343.282, 125.388),
    },
    "A1": { 
        "prop": (694.327, 258.500),
        "owner_main": (695.725, 252.500),
        "owner_aux": (786.257, 101.733),
        "doc_main": (684.368, 246.500),
        "doc_aux": (786.257, 96.551),
        "date": (673.691, 198.500),
        "area": (691.408, 115.228),
        "perimeter": (722.402, 115.228),
        "folha": (815.846, 286.482),
        "escala": (793.194, 198.500),
        "lat": (809.269, 513.357),
        "lon": (809.269, 509.357),
        "convergencia": (760.799, 517.357),
        "declinacao": (760.799, 513.357),
        "variacao": (760.799, 509.357),
        "municipio": (691.706, 240.500),
        "estado": (799.253, 240.500),
        "comarca": (691.706, 234.500),
        "cartorio": (686.244, 228.500),
        "cns": (673.021, 216.500),
    },
}

TABLE_HEADERS = [
    "Vértice-Vante",
    "E (m)",
    "N (m)",
    "h (m)",
    "Azimute",
    "Distancia (m)",
]

DEFAULT_COLUMN_WIDTHS = [20.0, 20.0, 20.0, 15.0, 20.0, 20.0]

DEFAULT_MISSING_TEXT = "Não informado"
ROAD_INTERSECTION_TOLERANCE = 1.0

# ---------------------------------------------------------------------------
# ESTRUTURAS DE DADOS
# ---------------------------------------------------------------------------


@dataclass
class SegmentRecord:
    start_vid: int
    end_vid: int
    start_point: Tuple[float, float]
    end_point: Tuple[float, float]
    start_label: str
    end_label: str
    azimute_txt: str
    distancia_txt: str
    azimute_value: float
    distancia_value: float
    h_value: Optional[float]
    h_txt: str


@dataclass
class PropertyInfo:
    label: str
    full_name: Optional[str]
    owner: Optional[str]
    document: Optional[str]
    area_ha: Optional[float]
    perimeter_m: Optional[float]
    municipio: Optional[str] = None
    uf: Optional[str] = None
    comarca: Optional[str] = None
    cartorio: Optional[str] = None


@dataclass
class ConfrontanteInfo:
    owner: Optional[str]
    label: Optional[str]
    from_road: bool = False


# ---------------------------------------------------------------------------
# HELPERS DE ORDENAMENTO E TEXTO
# ---------------------------------------------------------------------------


def _rotate_start_southwest(pts):
    if not pts:
        return pts
    idx = min(range(len(pts)), key=lambda i: (pts[i][1], pts[i][0]))
    return pts[idx:] + pts[:idx]


def _order_indices_by_chain_proximity(points: Sequence[Tuple[float, float]]) -> List[int]:
    """Return indices ordered by nearest-neighbour chaining starting from the most southern point."""
    if not points:
        return []
    start = min(range(len(points)), key=lambda i: (points[i][1], points[i][0]))
    remaining = set(range(len(points)))
    remaining.remove(start)
    ordered = [start]
    while remaining:
        last = ordered[-1]
        lx, ly = points[last]
        next_idx = min(
            remaining,
            key=lambda i: ((points[i][0] - lx) ** 2 + (points[i][1] - ly) ** 2, points[i][1], points[i][0]),
        )
        ordered.append(next_idx)
        remaining.remove(next_idx)
    return ordered


def _quantize_point(pt: Tuple[float, float], tol: float = 0.01) -> Tuple[int, int]:
    return (int(round(pt[0] / tol)), int(round(pt[1] / tol)))


def _normalize_ascii(text: str) -> str:
    return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")


def _normalize_whitespace(text: str) -> str:
    """Remove espaços duplicados e espaços extras no início/fim."""
    return re.sub(r"\s+", " ", text).strip()


def _clean_text_value(value: object) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, float) and math.isnan(value):
        return None
    text = str(value).strip()
    return text or None


def _parse_float(value: object) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, float):
        if math.isnan(value):
            return None
        return float(value)
    if isinstance(value, int):
        return float(value)
    text = str(value).strip()
    if not text:
        return None
    text = text.replace(" ", "")
    if text.count(",") == 1 and "." in text:
        text = text.replace(".", "")
    text = text.replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return None


def _label_sort_key(label: str, fallback: int) -> Tuple[int, object]:
    try:
        return (0, int(label))
    except (TypeError, ValueError):
        try:
            return (0, float(label))
        except (TypeError, ValueError):
            return (1, str(label), fallback)


def _sanitize_label(label: Optional[str]) -> Optional[str]:
    if label is None:
        return None
    lbl = label.strip().replace("_", " ")
    lbl = _normalize_ascii(lbl)
    return lbl or None


def _normalize_property_label(raw_label: Optional[str], fallback_idx: int) -> str:
    base = _clean_text_value(raw_label)
    if base:
        numeric = _parse_float(base)
        if numeric is not None:
            if numeric.is_integer():
                return f"{int(round(numeric)):02d}"
            text = format(numeric, "f").rstrip("0").rstrip(".")
            return text or f"{fallback_idx:02d}"
        sanitized = _sanitize_label(base)
        if sanitized:
            return sanitized
    return f"{fallback_idx:02d}"


def _safe_slug(text: str, fallback: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        return fallback
    if re.fullmatch(r"[0-9]+", cleaned):
        return cleaned
    try:
        value = float(cleaned)
    except ValueError:
        value = None
    else:
        if value.is_integer():
            return str(int(value))
    slug = re.sub(r"[^0-9A-Za-z_-]+", "_", cleaned)
    slug = slug.strip("_")
    return slug or fallback


def _coord_text(pt: Tuple[float, float]) -> str:
    return f"E = {_fmt_num(pt[0], 4)} m e N = {_fmt_num(pt[1], 4)} m"


def _coord_text_colon(pt: Tuple[float, float]) -> str:
    return f"E: {_fmt_num(pt[0], 4)} m e N: {_fmt_num(pt[1], 4)} m"


def _short_azimute(text: str) -> str:
    """Retorna azimute no formato GG°MM' (descarta segundos)."""
    if not text:
        return text
    m = re.match(r"\s*([0-9]+)°\s*([0-9]{1,2})", text)
    if m:
        deg, minutes = m.groups()
        return f"{deg}°{minutes}'"
    return text


def _load_municipios() -> Optional[gpd.GeoDataFrame]:
    paths = [p for p in [MUNICIPIOS_BA, MUNICIPIOS_SE] if p.exists()]
    if not paths:
        return None
    frames = []
    for p in paths:
        try:
            gdf = gpd.read_file(p)
            frames.append(gdf)
        except Exception as exc:
            _log_error(f"Falha ao ler municipios '{p}': {exc}")
    if not frames:
        return None
    try:
        gdf_all = gpd.pd.concat(frames, ignore_index=True)
    except Exception:
        return None
    cols = {c.lower(): c for c in gdf_all.columns}
    nm_col = cols.get("nm_mun") or cols.get("nome") or list(cols.values())[0]
    uf_col = cols.get("sigla_uf") or cols.get("uf")
    gdf_all = gdf_all.rename(columns={nm_col: "NM_MUN", uf_col: "SIGLA_UF"} if uf_col else {nm_col: "NM_MUN"})
    gdf_all = gdf_all[["NM_MUN", "SIGLA_UF", "geometry"]].copy()
    try:
        if gdf_all.crs is None or gdf_all.crs.to_string().upper() != GEOD_CRS:
            gdf_all = gdf_all.to_crs(GEOD_CRS)
    except Exception:
        pass
    return gdf_all


def _infer_municipio_for_poly(
    poly_vertices: Sequence[Tuple[float, float]],
    gdf_mun: Optional[gpd.GeoDataFrame],
) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]:
    if gdf_mun is None or not poly_vertices:
        return None, None, None, None
    try:
        poly = Polygon(poly_vertices)
    except Exception:
        return None, None, None, None
    try:
        centroid = poly.centroid
    except Exception:
        return None, None, None, None
    try:
        pt_gdf = gpd.GeoDataFrame(geometry=[centroid], crs=GEOD_CRS)
        joined = gpd.sjoin(pt_gdf, gdf_mun, how="left")
        if not joined.empty and not joined.iloc[0].isnull().all():
            mun = _clean_text_value(joined.iloc[0].get("NM_MUN"))
            uf = _clean_text_value(joined.iloc[0].get("SIGLA_UF"))
            comarca, cartorio = None, None
            if mun:
                mun_norm = mun.strip().lower()
                if mun_norm == "paulo afonso":
                    comarca = "Paulo Afonso"
                    cartorio = "Ofício de registro de imóveis títulos e documentos e pessoas jurídicas da Comarca de Paulo Afonso"
                elif mun_norm in {"canindé de são francisco", "caninde de são francisco"}:
                    comarca = "Canindé de São Francisco"
                    cartorio = "Cartório do Ofício Único da Comarca Canindé de São Francisco"
            return mun, uf, comarca, cartorio
    except Exception:
        return None, None, None, None
    return None, None, None, None


def _select_property_indices(poly_labels: Sequence[str], poly_info: Sequence[PropertyInfo]) -> Optional[List[int]]:
    """
    Interface simples para o usuário escolher quais propriedades exportar.
    Retorna lista de índices selecionados ou None para exportar todas.
    """
    if not SELECT_PROPERTIES_INTERACTIVE:
        return None
    try:
        print("\n== Seleção de propriedades ==")
        for i, label in enumerate(poly_labels, start=1):
            name = poly_info[i - 1].full_name if i - 1 < len(poly_info) else ""
            name_display = f" - {name}" if name else ""
            print(f"  {label}{name_display}")
        raw = input("Digite os IDs/nomes separados por vírgula (Enter para todas): ").strip()
    except Exception:
        return None
    if not raw:
        return None
    tokens = {t.strip().lower() for t in raw.split(",") if t.strip()}
    if not tokens:
        return None

    selected: List[int] = []
    for idx, label in enumerate(poly_labels):
        lbl = (label or "").strip().lower()
        name = (poly_info[idx].full_name if idx < len(poly_info) else "") or ""
        if lbl in tokens or name.strip().lower() in tokens:
            selected.append(idx)
    return selected or None


# ---------------------------------------------------------------------------
# VÉRTICES / CONFRONTANTES
# ---------------------------------------------------------------------------


class VertexIndex:
    """Deduplicação global de vértices aproximados."""

    def __init__(self, tol: float = 0.01) -> None:
        self.tol = float(tol)
        self._grid: Dict[Tuple[int, int], Tuple[int, Tuple[float, float]]] = {}
        self._next_id = 1
        self._drawn_ids: set[int] = set()

    def _key(self, x: float, y: float) -> Tuple[int, int]:
        return (int(round(x / self.tol)), int(round(y / self.tol)))

    def get_vid(self, x: float, y: float) -> int:
        k = self._key(x, y)
        if k in self._grid:
            return self._grid[k][0]
        vid = self._next_id
        self._next_id += 1
        self._grid[k] = (vid, (x, y))
        return vid

    def should_draw(self, vid: int) -> bool:
        return vid not in self._drawn_ids

    def mark_drawn(self, vid: int) -> None:
        self._drawn_ids.add(vid)


def _segment_key(rec: SegmentRecord, tol: float = 0.01) -> Tuple:
    pts = tuple(sorted((_quantize_point(rec.start_point, tol), _quantize_point(rec.end_point, tol))))
    return pts


def _build_confrontantes_map(
    poly_records: Sequence[Sequence[SegmentRecord]],
    poly_labels: Sequence[str],
    poly_info: Sequence[PropertyInfo],
    road_geometries: Sequence[BaseGeometry],
    road_tolerance: float,
) -> Dict[Tuple[int, int], ConfrontanteInfo]:
    edge_map: Dict[Tuple, List[Tuple[int, int]]] = {}
    for poly_idx, records in enumerate(poly_records):
        for rec_idx, rec in enumerate(records):
            key = _segment_key(rec)
            edge_map.setdefault(key, []).append((poly_idx, rec_idx))

    confrontantes: Dict[Tuple[int, int], ConfrontanteInfo] = {}
    for matches in edge_map.values():
        if len(matches) <= 1:
            continue
        involved_polys = {poly_idx for poly_idx, _ in matches}
        for poly_idx, rec_idx in matches:
            neighbor_idx = next((idx for idx in involved_polys if idx != poly_idx), None)
            if neighbor_idx is None:
                continue
            owner = None
            label = None
            if 0 <= neighbor_idx < len(poly_info):
                owner = _clean_text_value(poly_info[neighbor_idx].owner)
                label = poly_info[neighbor_idx].label
            if label is None and 0 <= neighbor_idx < len(poly_labels):
                label = poly_labels[neighbor_idx]
            confrontantes[(poly_idx, rec_idx)] = ConfrontanteInfo(owner=owner, label=label, from_road=False)

    if road_geometries:
        tolerance = max(road_tolerance, 0.0)
        buffered_cache: Dict[Tuple[int, int], BaseGeometry] = {}
        for poly_idx, records in enumerate(poly_records):
            for rec_idx, rec in enumerate(records):
                key = (poly_idx, rec_idx)
                info = confrontantes.get(key)
                if info is not None and not info.from_road:
                    continue

                segment_line = LineString([rec.start_point, rec.end_point])
                if tolerance > 0.0:
                    cache_key = (poly_idx, rec_idx)
                    if cache_key not in buffered_cache:
                        buffered_cache[cache_key] = segment_line.buffer(tolerance, cap_style=2)
                    segment_geom = buffered_cache[cache_key]
                else:
                    segment_geom = segment_line

                for geom in road_geometries:
                    if geom is None or geom.is_empty:
                        continue
                    try:
                        if segment_geom.intersects(geom):
                            confrontantes[key] = ConfrontanteInfo(owner=None, label=None, from_road=True)
                            break
                    except Exception:
                        continue
    return confrontantes


def _confrontante_text(info: Optional[ConfrontanteInfo]) -> str:
    if info is not None:
        if info.from_road:
            return "Estrada Municipal"

        owner_clean = _normalize_whitespace(info.owner) if info.owner else None
        if owner_clean:
            return f"propriedade de {owner_clean}"

        label = _sanitize_label(info.label)
        if label:
            normalized = _normalize_ascii(label).lower()
            if normalized in {"ada", "area de atuacao (ada)", "area de atuacao ada", "area de atuacao"}:
                return "GLEBA"

            return f"propriedade {label}"

    return "GLEBA"


# ---------------------------------------------------------------------------
# FORMATOS / NÚMEROS / GEOMETRIA
# ---------------------------------------------------------------------------


def _fmt_pid(vid: int) -> str:
    return f"P{vid:02d}"


def _fmt_num(value: Optional[float], ndigits: int = 2) -> str:
    if value is None:
        return "-"
    s = f"{value:.{ndigits}f}"
    return s.replace(".", ",")


def _deg_to_dms(deg: float) -> Tuple[int, int, float]:
    d = int(deg)
    m_f = abs(deg - d) * 60.0
    m = int(m_f)
    s = (m_f - m) * 60.0
    return d, m, abs(s)


def _dms_str(az_deg: float) -> str:
    d, m, s = _deg_to_dms(az_deg % 360.0)
    deg_sym = chr(176)
    seconds = f"{s:05.2f}".replace(".", ",")
    return f"{d:02d}{deg_sym}{m:02d}'{seconds}\""


def _dms_str_dim(az_deg: float) -> str:
    d, m, s = _deg_to_dms(az_deg % 360.0)
    deg_sym = chr(176)
    return f"{d:02d}{deg_sym}{m:02d}'{s:02.0f}\""


def _dist(a: Tuple[float, float], b: Tuple[float, float]) -> float:
    return float(math.hypot(b[0] - a[0], b[1] - a[1]))


def _azimute(a: Tuple[float, float], b: Tuple[float, float]) -> float:
    dE = b[0] - a[0]
    dN = b[1] - a[1]
    ang = math.degrees(math.atan2(dE, dN))
    return ang if ang >= 0 else ang + 360.0


def _bulge_to_delta_deg(bulge: float) -> float:
    return math.degrees(4.0 * math.atan(bulge))


def _radius_from_chord(chord_len: float, delta_deg: float) -> Optional[float]:
    half = math.radians(abs(delta_deg) / 2.0)
    s = math.sin(half)
    if abs(s) < 1e-9:
        return None
    return chord_len / (2.0 * s)


def _poly_bbox(vertices: Sequence[Tuple[float, float]]) -> Tuple[float, float, float, float]:
    xs = [p[0] for p in vertices]
    ys = [p[1] for p in vertices]
    return min(xs), min(ys), max(xs), max(ys)


def _polygon_centroid(vertices: Sequence[Tuple[float, float]]) -> Tuple[float, float]:
    if not vertices:
        return 0.0, 0.0
    area = 0.0
    cx = 0.0
    cy = 0.0
    n = len(vertices)
    for i in range(n):
        x1, y1 = vertices[i]
        x2, y2 = vertices[(i + 1) % n]
        cross = x1 * y2 - x2 * y1
        area += cross
        cx += (x1 + x2) * cross
        cy += (y1 + y2) * cross
    if abs(area) < 1e-9:
        return (sum(x for x, _ in vertices) / n, sum(y for _, y in vertices) / n)
    area *= 0.5
    cx /= 6.0 * area
    cy /= 6.0 * area
    return cx, cy


def _add_polyline_with_bulge(
    ms,
    points: Sequence[Tuple[float, float, float]],
    layer: str,
    color: Optional[int] = None,
):
    if not points:
        return None
    data = [(x, y, 0.0, 0.0, bulge) for x, y, bulge in points]
    attribs = {"layer": layer}
    if color is not None:
        attribs["color"] = color
    return ms.add_lwpolyline(data, format="xyseb", close=True, dxfattribs=attribs)


def _add_metadata_text(
    doc: ezdxf.EzDxfDocument,
    layout: Optional[ezdxf.layouts.BaseLayout],
    text: str,
    x: float,
    y: float,
    params: Params,
    layer: str,
    align: str = "MIDDLE_LEFT",
    height: float = 3.0,
) -> None:
    if layout is None:
        return
    value = text if text else DEFAULT_MISSING_TEXT
    style_name = params.style_texto
    if style_name not in doc.styles:
        try:
            doc.styles.add(style_name, font="arial.ttf")
        except Exception:
            style_name = "Standard"
    try:
        _add_text_align(
            layout,
            value,
            x,
            y,
            height,
            layer,
            align=align,
            style=style_name,
        )
    except Exception as exc:
        _log_error(
            f"Nao foi possivel inserir texto '{value}' no layout '{getattr(layout, 'name', 'desconhecido')}'. "
            f"Motivo: {exc}"
        )


def _write_layout_metadata(
    doc: ezdxf.EzDxfDocument,
    layout: Optional[ezdxf.layouts.BaseLayout],
    params: Params,
    metadata: Sequence[Tuple[str, float, float] | Tuple[str, float, float, str] | Tuple[str, float, float, str, float]],
    folha_value: str,
    layer: str,
) -> None:
    if layout is None:
        _log_error("Layout ausente para insercao de metadados do memorial.")
        return
    if layer not in doc.layers:
        try:
            doc.layers.add(layer)
        except Exception:
            pass
    positions = _carimbo_positions(layout)
    for entry in metadata:
        if len(entry) == 5:
            text, x, y, align, height = entry
        elif len(entry) == 4:
            text, x, y, align = entry
            height = 3.0
        else:
            text, x, y = entry  # type: ignore[misc]
            align = "MIDDLE_LEFT"
            height = 3.0
        _add_metadata_text(doc, layout, text, x, y, params, layer, align=align, height=height)
    folha_x, folha_y = positions.get("folha", (43.129, 29.700))
    _add_metadata_text(doc, layout, folha_value, folha_x, folha_y, params, layer)


def _debug(msg: str) -> None:
    if DEBUG_MODE:
        print(f"[DEBUG] {msg}")


def _log_error(msg: str) -> None:
    print(f"[ERRO] {msg}")


def _find_layout_by_name(doc: ezdxf.EzDxfDocument, name: str) -> Optional[ezdxf.layouts.BaseLayout]:
    try:
        layout = doc.layout(name)  # type: ignore[call-arg, attr-defined]
        if layout is not None:
            return layout
    except Exception:
        pass
    try:
        layout = doc.layouts.get_layout_by_name(name)  # type: ignore[attr-defined]
        return layout
    except Exception:
        return None

def _get_preferred_layout(doc: ezdxf.EzDxfDocument) -> Optional[ezdxf.layouts.BaseLayout]:
    # Forçar uso do layout A2; se não existir, cria e usa.
    layout = _find_layout_by_name(doc, "A2")
    if layout is not None:
        return layout
    try:
        layout = doc.layouts.new("A2")
        return layout
    except Exception:
        return None


def _carimbo_positions(layout: Optional[ezdxf.layouts.BaseLayout]) -> Dict[str, Tuple[float, float]]:
    """Seleciona conjunto de posições do carimbo conforme o nome do layout (A1/A2/A3)."""
    name = getattr(layout, "name", "") or ""
    key = name.upper()
    if key in CARIMBO_POSITIONS:
        return CARIMBO_POSITIONS[key]
    return CARIMBO_POSITIONS[CARIMBO_DEFAULT_LAYOUT]


def _table_anchor_for_layout(layout: Optional[ezdxf.layouts.BaseLayout]) -> Tuple[float, float]:
    """Retorna a âncora da tabela conforme o layout, caindo no padrão se não encontrado."""
    name = getattr(layout, "name", "") or ""
    anchor = TABLE_LAYOUT_ANCHORS.get(name.upper())
    if anchor is not None:
        return anchor
    return (TABLE_ANCHOR_NO_ARC_X, TABLE_ANCHOR_NO_ARC_Y)


def _cartorio_for_layout(cartorio_txt: str, comarca_txt: str, layout: Optional[ezdxf.layouts.BaseLayout]) -> str:
    """
    Ajusta o texto do cartório conforme o layout:
    - Mantém o texto do cartório separado do texto de comarca (não agrega nem corta a comarca).
    """
    return cartorio_txt


def _cartorio_metadata_entries(
    layout: Optional[ezdxf.layouts.BaseLayout],
    positions: Dict[str, Tuple[float, float]],
    cartorio_txt: str,
    comarca_txt: str,
    municipio_txt: str,
    cartorio_display: str,
) -> Tuple[Tuple[str, float, float, str, float], List[Tuple[str, float, float, str, float]]]:
    """
    Define entradas de metadados de comarca/cart¢rio com tratamentos especiais
    para a folha A3 quando as comarcas sÆo Paulo Afonso ou Canind‚ de SÆo Francisco.
    A comarca SEMPRE usa as coordenadas do layout (CARIMBO_POSITIONS) para nao
    se confundir com o split do cartorio.
    """
    layout_name = (getattr(layout, "name", "") or "").upper()

    norm_cartorio = _normalize_ascii(cartorio_txt).upper()
    norm_comarca = _normalize_ascii(comarca_txt).upper()
    norm_municipio = _normalize_ascii(municipio_txt).upper()

    # Comarca exibe apenas o nome do munic¡pio (sem prefixar "Comarca de ...")
    comarca_label = municipio_txt or comarca_txt

    # MantÉm o cart¢rio completo (sem cortar sufixos)
    base_cartorio_display = cartorio_display

    default_comarca = (comarca_label, *positions.get("comarca", (0.0, 0.0)), "MIDDLE_LEFT", 3.0)
    default_cartorio = [(base_cartorio_display, *positions.get("cartorio", (0.0, 0.0)), "MIDDLE_LEFT", 3.0)]

    if layout_name != "A3":
        return default_comarca, default_cartorio

    if "PAULO AFONSO" in norm_cartorio or "PAULO AFONSO" in norm_comarca:
        comarca_entry = default_comarca
        cartorio_entries = [
            ("Ofício de registro de imóveis", 356.505, 143.388, "MIDDLE_LEFT", 3.0),
            ("títulos e documentos e pessoas jurídicas da", 329.260, 137.388, "MIDDLE_LEFT", 3.0),
        ]
        return comarca_entry, cartorio_entries

    if "CANINDE DE SAO FRANCISCO" in norm_cartorio or "CANINDE DE SAO FRANCISCO" in norm_comarca:
        comarca_entry = default_comarca
        cartorio_entries = [
            ((cartorio_txt or base_cartorio_display), *positions.get("cartorio", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
        ]
        return comarca_entry, cartorio_entries

    return default_comarca, default_cartorio


def _get_viewport_scale(layout: Optional[ezdxf.layouts.BaseLayout], layer_name: str = "CAD_VIEWPORT") -> Optional[str]:
    if layout is None:
        return None
    for e in layout:
        try:
            if e.dxftype() == "VIEWPORT" and e.dxf.layer == layer_name:
                vp = e  # type: ignore[assignment]
                vp_h = float(getattr(vp.dxf, "height", 0.0) or 0.0)
                vh = float(getattr(vp.dxf, "view_height", 0.0) or 0.0)
                if vp_h > 0.0 and vh > 0.0:
                    denom = vh * 1000.0 / vp_h
                    return f"1/{int(round(denom))}"
        except Exception:
            continue
    return None


def _adjust_viewports_for_property(
    layout: Optional[ezdxf.layouts.BaseLayout],
    bbox: Tuple[float, float, float, float],
) -> None:
    """
    Centraliza a visão nas viewports com base no bbox da propriedade.

    - CAD_VIEWPORT: zoom extents automático (preenche o quadro).
    - CAD_SITUACAO / CAD_LOCALIZACAO: zoom extents e depois escala fixa 1:30.000.
    """
    if layout is None:
        return

    minx, miny, maxx, maxy = bbox
    bw = maxx - minx
    bh = maxy - miny
    if bw <= 0 or bh <= 0:
        return

    cx = (minx + maxx) / 2.0
    cy = (miny + maxy) / 2.0

    def _zoom_extents(
        vp: ezdxf.entities.Viewport,
        center: Tuple[float, float],
        bbox_size: Tuple[float, float],
        padding: float = 1.10,
    ) -> Optional[float]:
        """Zoom extents dentro da viewport (retorna view_height calculado)."""
        vp_w = float(getattr(vp.dxf, "width", 0.0) or 0.0)
        vp_h = float(getattr(vp.dxf, "height", 0.0) or 0.0)
        if vp_w <= 0.0 or vp_h <= 0.0:
            return None

        bw_l, bh_l = bbox_size
        vh_por_altura = bh_l * padding
        vh_por_largura = bw_l * padding * (vp_h / vp_w)
        view_height = max(vh_por_altura, vh_por_largura)

        # Ajuste para escalas normalizadas em passos de 1/500 até 1/10000
        try:
            denom = view_height * 1000.0 / vp_h  # escala 1/denom
            snapped = min(max(math.ceil(denom / 500.0) * 500.0, 500.0), 10000.0)
            view_height = snapped * vp_h / 1000.0
        except Exception:
            pass

        vp.dxf.view_height = view_height

        try:
            vp.dxf.view_target_point = (center[0], center[1], 0.0)
        except Exception:
            pass
        try:
            vp.dxf.view_direction_vector = (0.0, 0.0, 1.0)
            vp.dxf.view_twist_angle = 0.0
        except Exception:
            pass
        try:
            vp.dxf.view_center_point = (0.0, 0.0, 0.0)
        except Exception:
            pass

        return view_height

    def _zoom_extents_and_fix_scale(
        vp: ezdxf.entities.Viewport,
        center: Tuple[float, float],
        bbox_size: Tuple[float, float],
        escala: float = 50000.0,
        padding: float = 1.05,
    ) -> None:
        """Zoom extents + aplica escala fixa (1:escala)."""
        vp_w = float(getattr(vp.dxf, "width", 0.0) or 0.0)
        vp_h = float(getattr(vp.dxf, "height", 0.0) or 0.0)
        if vp_w <= 0.0 or vp_h <= 0.0:
            return

        _zoom_extents(vp, center, bbox_size, padding=padding)

        view_height_fixed = escala * (vp_h / 1000.0)
        vp.dxf.view_height = view_height_fixed
        try:
            vp.dxf.view_target_point = (center[0], center[1], 0.0)
        except Exception:
            pass
        try:
            vp.dxf.view_direction_vector = (0.0, 0.0, 1.0)
            vp.dxf.view_twist_angle = 0.0
        except Exception:
            pass
        try:
            vp.dxf.view_center_point = (0.0, 0.0, 0.0)
        except Exception:
            pass
        try:
            vp.frozen_layers.add(LAYER_VERTICE_PTO)
            vp.frozen_layers.add(LAYER_VERTICE_TXT)
        except Exception:
            pass

    def _find_vp(layer_name: str) -> Optional[ezdxf.entities.Viewport]:
        for e in layout:
            try:
                if e.dxftype() == "VIEWPORT" and e.dxf.layer == layer_name:
                    return e  # type: ignore[return-value]
            except Exception:
                continue
        return None

    vp_main = _find_vp("CAD_VIEWPORT")
    if vp_main is None:
        _debug(f"[VIEWPORT] Nenhuma viewport principal (CAD_VIEWPORT) encontrada no layout '{getattr(layout, 'name', 'desconhecido')}'.")
    if vp_main is not None:
        _zoom_extents(vp_main, (cx, cy), (bw, bh), padding=1.10)

    vp_sit = _find_vp("CAD_SITUACAO")
    if vp_sit is None:
        _debug(f"[VIEWPORT] Nenhuma viewport de situacao (CAD_SITUACAO) encontrada no layout '{getattr(layout, 'name', 'desconhecido')}'.")
    if vp_sit is not None:
        _zoom_extents_and_fix_scale(vp_sit, (cx, cy), (bw, bh), escala=30000.0, padding=1.05)

    vp_loc = _find_vp("CAD_LOCALIZACAO")
    if vp_loc is None:
        _debug(f"[VIEWPORT] Nenhuma viewport de localizacao (CAD_LOCALIZACAO) encontrada no layout '{getattr(layout, 'name', 'desconhecido')}'.")
    if vp_loc is not None:
        _zoom_extents_and_fix_scale(vp_loc, (cx, cy), (bw, bh), escala=30000.0, padding=1.05)


def _extract_id_and_name(raw: Optional[str]) -> Tuple[Optional[int], Optional[str], Optional[str]]:
    if not raw:
        return None, None, None
    text = str(raw).strip()
    if not text:
        return None, None, None
    m = re.match(r"\s*(\d+)([A-Za-z]?)\s*[--]?\s*(.*)", text)
    if m:
        num = int(m.group(1))
        suffix = m.group(2).strip().upper() or None
        name = m.group(3).strip() or None
        return num, suffix, name
    return None, None, text


def _make_height_sampler(mdt_path: Path):
    if not mdt_path.exists():
        _debug(f"MDT nao encontrado em {mdt_path}")
        return None, None
    try:
        ds = rasterio.open(mdt_path)
    except Exception as exc:
        _log_error(f"Falha ao abrir MDT '{mdt_path}': {exc}")
        return None, None

    nodata = ds.nodata

    def sampler(pt: Tuple[float, float]) -> Optional[float]:
        x, y = pt
        try:
            val = next(ds.sample([(x, y)]))[0]
        except Exception:
            return None
        try:
            v = float(val)
        except Exception:
            return None
        if nodata is not None and abs(v - nodata) < 1e-6:
            return None
        if math.isnan(v):
            return None
        return v

    return ds, sampler


def _format_latlon(value: float, is_lat: bool) -> str:
    txt = f"{value:.4f}".replace(".", ",")
    return f"{txt}°"


def _compute_geo_info(x: float, y: float) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str], Optional[str]]:
    """
    Converte XY (EPSG 31984) para lat/lon, calcula convergência meridiana
    e consulta declinação/variação magnética no serviço NOAA (geomag).
    """
    try:
        from pyproj import Transformer
    except Exception:
        _log_error("pyproj não disponível para converter coordenadas.")
        return None, None, None, None, None

    try:
        transformer = Transformer.from_crs(GEOD_CRS, "EPSG:4674", always_xy=True)
        lon, lat = transformer.transform(x, y)
    except Exception as exc:
        _log_error(f"Falha ao converter coordenadas para lat/lon: {exc}")
        return None, None, None, None, None

    lon0 = -39.0  # Meridiano central UTM 24S (EPSG:31984)
    lat_rad = radians(lat)
    dlon_rad = radians(lon - lon0)
    try:
        conv_rad = atan(tan(dlon_rad) * sin(lat_rad))
        conv_deg = degrees(conv_rad)
        conv_txt = f"{conv_deg:.3f}°"
    except Exception:
        conv_txt = None

    lat_txt = _format_latlon(lat, True)
    lon_txt = _format_latlon(lon, False)

    decl_txt, var_txt = None, None
    today = datetime.now()
    # 1) tenta geomag local (offline) - pacote geomag (retorna apenas declinação)
    if geomag is not None:
        try:
            t_ref = datetime(today.year, 12, min(today.day, 28)).date()  # usa dezembro do ano atual (date)
            dec_val = geomag.declination(lat, lon, h=0.0, time=t_ref)
            if dec_val is not None:
                decl_txt = f"{float(dec_val):.3f}°"
        except Exception as exc:
            _log_error(f"Falha na declinação via geomag: {exc}")

    if var_txt is None:
        var_txt = "0,10°"
    return lat_txt, lon_txt, conv_txt, decl_txt, var_txt


def _extract_polygons_from_shapefile(path: Path) -> Optional[Tuple[List[List[Tuple[float, float, float]]], List[str], List[PropertyInfo], List[Optional[int]], List[BaseGeometry]]]:
    """Lê PROPRIEDADES.shp e devolve geometrias, labels, info e geometria de estrada."""
    if not path.exists():
        return None
    try:
        gdf = gpd.read_file(path)
    except Exception as exc:
        _log_error(f"Falha ao ler shapefile de propriedades '{path}': {exc}")
        return None
    if gdf.empty:
        _log_error(f"Shapefile '{path}' está vazio.")
        return None

    cols = {c.lower(): c for c in gdf.columns}
    nm_col = cols.get("nm_propr")
    owner_col = cols.get("prop_nome")
    doc_col = cols.get("pro_cpf") or cols.get("prop_cpf")
    area_col = cols.get("area_1")
    perim_col = cols.get("perimetr_1")

    poly_pts: List[List[Tuple[float, float, float]]] = []
    labels: List[str] = []
    infos: List[PropertyInfo] = []
    id_values: List[Optional[int]] = []
    road_geoms: List[BaseGeometry] = []

    for _, row in gdf.iterrows():
        geom = row.geometry
        if geom is None or geom.is_empty:
            continue

        name_raw = _clean_text_value(row.get(nm_col)) if nm_col else None
        name_ascii = _normalize_ascii(name_raw) if name_raw else ""
        if name_ascii and "ESTRADA" in name_ascii.upper():
            if geom.geom_type in {"Polygon", "MultiPolygon"}:
                g = geom.boundary
            else:
                g = geom
            if g is not None and not g.is_empty:
                road_geoms.append(g)
            continue

        geom_use = geom
        if geom.geom_type == "MultiPolygon":
            parts = list(geom.geoms)
            if not parts:
                continue
            geom_use = max(parts, key=lambda p: p.area)

        if geom_use.geom_type != "Polygon":
            geom_use = geom_use.convex_hull

        coords = list(getattr(geom_use, "exterior", geom_use).coords)
        pts = [(float(x), float(y), 0.0) for x, y in coords]
        if len(pts) < 4:  # precisa fechar
            continue

        num_id, suffix_id, name_clean = _extract_id_and_name(name_raw)
        if num_id is not None:
            if suffix_id:
                label_text = f"{num_id:02d}{suffix_id}"
            else:
                label_text = f"{num_id:02d}"
        else:
            label_text = name_clean or name_raw or f"{len(poly_pts)+1:02d}"

        owner_val = _clean_text_value(row.get(owner_col)) if owner_col else None
        doc_val = _clean_text_value(row.get(doc_col)) if doc_col else None
        area_val = _parse_float(row.get(area_col)) if area_col else None
        per_val = _parse_float(row.get(perim_col)) if perim_col else None

        poly_pts.append(pts)
        labels.append(label_text)
        infos.append(
            PropertyInfo(
                label=label_text,
                full_name=name_raw,
                owner=owner_val,
                document=doc_val,
                area_ha=area_val,
                perimeter_m=per_val,
                municipio=None,
                uf=None,
                comarca=None,
                cartorio=None,
            )
        )
        id_values.append(num_id)

    if not poly_pts:
        return None

    return poly_pts, labels, infos, id_values, road_geoms


# ---------------------------------------------------------------------------
# PARÂMETROS DE DESENHO
# ---------------------------------------------------------------------------


@dataclass
class Params:
    altura_texto_P: float = 3.0
    p_label_dx: float = 1.18
    p_label_dy: float = 0.34
    text_gap: float = 0.0
    altura_texto_tabela: float = 2.0
    tabela_cell_w: float = 25.0
    tabela_cell_h: float = 6.0
    tabela_offset_x: float = 120.0
    tabela_offset_y: float = 0.0
    altura_texto_ordinate: float = 0.9
    ord_dx: float = 4.0
    ord_dy: float = 4.0
    dimtxt_ordinate: float = 0.9
    dimasz_ordinate: float = 0.25
    style_texto: str = STYLE_TEXTO


# ---------------------------------------------------------------------------
# DXF: LEITURA DE POLYLINES
# ---------------------------------------------------------------------------


def _iter_closed_polylines(
    msp: ezdxf.layouts.BaseLayout, layers: Optional[Sequence[str]] = None
) -> Iterable[Tuple[str, List[Tuple[float, float, float]]]]:
    allowed = set(layers) if layers else None

    def layer_ok(name: str) -> bool:
        return True if allowed is None else name in allowed

    for e in msp.query("LWPOLYLINE"):
        if not e.closed or not layer_ok(e.dxf.layer):
            continue
        pts = [(x, y, bulge) for x, y, _, _, bulge in e.get_points()]
        yield e.dxf.layer, pts

    for e in msp.query("POLYLINE"):
        if e.is_3d_polyline or e.is_polygon_mesh or e.is_poly_face_mesh:
            continue
        if not e.is_closed or not layer_ok(e.dxf.layer):
            continue
        pts = []
        for v in e.vertices:
            bulge = getattr(v.dxf, "bulge", 0.0)
            pts.append((v.dxf.location.x, v.dxf.location.y, bulge))
        yield e.dxf.layer, pts


# ---------------------------------------------------------------------------
# DXF: DESENHO DE TEXTO / TABELAS / VÉRTICES
# ---------------------------------------------------------------------------


def _add_text_align(ms, text, cx, cy, h, layer, align="MIDDLE_CENTER", style=STYLE_TEXTO):
    t = ms.add_text(text, dxfattribs={"height": h, "style": style, "layer": layer})
    try:
        t.set_pos((cx, cy), align=align)
    except AttributeError:
        halign_map = {"LEFT": 0, "CENTER": 1, "RIGHT": 2}
        valign_map = {"BASELINE": 0, "BOTTOM": 1, "MIDDLE": 2, "TOP": 3}
        parts = align.split("_")
        v, h_align = ("MIDDLE", "CENTER")
        if len(parts) == 2:
            v, h_align = parts[0], parts[1]
        t.dxf.insert = (cx, cy)
        t.dxf.align_point = (cx, cy)
        t.dxf.halign = halign_map.get(h_align, 1)
        t.dxf.valign = valign_map.get(v, 2)
    return t


def _add_text_center(ms, text, cx, cy, h, layer, style=STYLE_TEXTO):
    return _add_text_align(ms, text, cx, cy, h, layer, "MIDDLE_CENTER", style)


def _draw_cell(ms, x, y, w, h, layer):
    ms.add_lwpolyline(
        [(x, y), (x + w, y), (x + w, y - h), (x, y - h), (x, y)],
        close=True,
        dxfattribs={"layer": layer},
    )


def _table_header(ms, x0, y0, widths: Sequence[float], ch: float, txt_h: float, layer: str, style: str, columns: Sequence[str]):
    if not widths:
        widths = [25.0] * len(columns)
    x = x0
    for idx, header in enumerate(columns):
        width = widths[idx] if idx < len(widths) else widths[-1]
        _draw_cell(ms, x, y0, width, ch, layer)
        _add_text_center(ms, header, x + width / 2.0, y0 - ch / 2.0, txt_h, layer, style)
        x += width


def _table_row(
    ms,
    row_idx: int,
    x0: float,
    y0: float,
    widths: Sequence[float],
    ch: float,
    txt_h: float,
    layer: str,
    style: str,
    values: Sequence[str],
    header_rows: int = 1,
):
    if not widths:
        widths = [25.0] * len(values)
    y = y0 - (header_rows + row_idx) * ch
    x = x0
    for idx, val in enumerate(values):
        width = widths[idx] if idx < len(widths) else widths[-1]
        _draw_cell(ms, x, y, width, ch, layer)
        _add_text_center(ms, val, x + width / 2.0, y - ch / 2.0, txt_h, layer, style)
        x += width


def _render_table(
    layout,
    params: Params,
    headers: Sequence[str],
    column_widths: Sequence[float],
    table_layer: str,
    style_texto: str,
    table_records: Sequence[SegmentRecord],
    x0: float,
    y0: float,
) -> None:
    ch = params.tabela_cell_h
    th = params.altura_texto_tabela
    column_widths = list(column_widths)
    total_width = sum(column_widths) if column_widths else 25.0 * max(len(headers), 1)

    # Linhas extras acima do cabeçalho com informações do sistema/projeção
    info_line1 = "Sistema de Referência - SIRGAS 2000"
    info_line2 = "Projeção UTM - FUSO 24 SUL (MC = -39º Wgr)"
    _draw_cell(layout, x0, y0, total_width, ch, table_layer)
    _add_text_center(layout, info_line1, x0 + total_width / 2.0, y0 - ch / 2.0, th, table_layer, style_texto)

    y_info2 = y0 - ch
    _draw_cell(layout, x0, y_info2, total_width, ch, table_layer)
    _add_text_center(layout, info_line2, x0 + total_width / 2.0, y_info2 - ch / 2.0, th, table_layer, style_texto)

    header_y = y0 - 2 * ch
    _table_header(layout, x0, header_y, column_widths, ch, th, table_layer, style_texto, headers)

    for row_idx, rec in enumerate(table_records):
        values = [
            f"{rec.start_label}-{rec.end_label}",
            _fmt_num(rec.start_point[0], 2),
            _fmt_num(rec.start_point[1], 2),
            rec.h_txt,
            rec.azimute_txt,
            rec.distancia_txt,
        ]
        _table_row(
            layout,
            row_idx,
            x0,
            y0,
            column_widths,
            ch,
            th,
            table_layer,
            style_texto,
            values,
            header_rows=3,
        )


def _draw_table_only(
    layout,
    doc: ezdxf.EzDxfDocument,
    params: Params,
    headers: Sequence[str],
    column_widths: Sequence[float],
    table_layer: str,
    records: Sequence[SegmentRecord],
    x0: float,
    y0: float,
    property_label: Optional[str] = None,
) -> None:
    if layout is None:
        _log_error(
            f"Nao foi possivel inserir a tabela secundaria para '{property_label or 'propriedade'}' "
            "porque nenhum layout valido foi encontrado."
        )
        return
    if not records:
        return
    if table_layer not in doc.layers:
        doc.layers.add(table_layer)
    _render_table(
        layout,
        params,
        headers,
        column_widths,
        table_layer,
        params.style_texto,
        records,
        x0,
        y0,
    )


def _register_vertex_block(doc: ezdxf.EzDxfDocument, radius: float = VERTEX_RADIUS):
    if BLOCO_VERTICE in doc.blocks:
        return
    blk = doc.blocks.new(BLOCO_VERTICE)
    blk.add_circle((0, 0), radius)
    # Cross (X) centered at origin, 3.5 x 3.5
    half = 3.5 / 2.0
    blk.add_line((-half, -half), (half, half))
    blk.add_line((-half, half), (half, -half))


def _register_label_block(
    doc: ezdxf.EzDxfDocument,
    dx: float,
    dy: float,
    text_h: float,
    style: str,
    layer_text: str,
    layer_leader: Optional[str] = None,
    name: str = "LBL_VERTICE",
) -> str:
    if name in doc.blocks:
        return name
    blk = doc.blocks.new(name)
    if layer_leader:
        blk.add_lwpolyline([(0.0, 0.0), (dx, dy)], dxfattribs={"layer": layer_leader})
    att = blk.add_attdef(
        tag="PID",
        text="P00",
        insert=(dx, dy),
        dxfattribs={"height": text_h, "style": style, "layer": layer_text},
    )
    att.dxf.halign = 0
    att.dxf.valign = 2
    return name


def _add_basic_layers(doc: ezdxf.EzDxfDocument):
    for layer in [LAYER_TABELA, LAYER_VERTICE_PTO, LAYER_VERTICE_TXT, LAYER_ORDINATE, CARIMBO_LAYER, "HACHURA"]:
        if layer not in doc.layers:
            doc.layers.add(layer)
    if STYLE_TEXTO != "Standard" and STYLE_TEXTO not in doc.styles:
        try:
            doc.styles.add(STYLE_TEXTO, font="txt")
        except Exception:
            pass


def _register_all_blocks(doc: ezdxf.EzDxfDocument, params: Params) -> str:
    _register_vertex_block(doc, radius=VERTEX_RADIUS)
    label_block = _register_label_block(
        doc,
        dx=-abs(params.p_label_dx),
        dy=abs(params.p_label_dy),
        text_h=params.altura_texto_P,
        style=params.style_texto,
        layer_text=LAYER_VERTICE_TXT,
    )
    return label_block


def _table_dims(
    n_rows: int,
    widths: Optional[Sequence[float]],
    ch: float,
    include_title: bool = True,
    header_rows: int = 1,
) -> Tuple[float, float]:
    if widths:
        width = sum(widths)
    else:
        width = sum(DEFAULT_COLUMN_WIDTHS)
    base_h = (header_rows + n_rows) * ch
    title_extra = 0.9 * ch if include_title else 0.0
    margin = 0.2 * ch
    height = base_h + title_extra + margin
    return width, height


# ---------------------------------------------------------------------------
# CONSTRUÇÃO DOS SEGMENTOS (RETA / ARCO)
# ---------------------------------------------------------------------------


def _build_segment_records(
    pts: List[Tuple[float, float, float]],
    vindex: Optional[VertexIndex],
    height_sampler: Optional[callable] = None,
) -> Tuple[List[Tuple[float, float]], List[int], List[SegmentRecord]]:
    # Remove vértices consecutivos quase coincidentes (< VERTEX_MIN_SPACING)
    cleaned_pts: List[Tuple[float, float, float]] = []
    for x, y, b in pts:
        if not cleaned_pts:
            cleaned_pts.append((x, y, b))
            continue
        lx, ly, _ = cleaned_pts[-1]
        if _dist((x, y), (lx, ly)) >= VERTEX_MIN_SPACING:
            cleaned_pts.append((x, y, b))

    # Se o último ficou colado ao primeiro, remove o último
    if len(cleaned_pts) >= 2 and _dist((cleaned_pts[-1][0], cleaned_pts[-1][1]), (cleaned_pts[0][0], cleaned_pts[0][1])) < VERTEX_MIN_SPACING:
        cleaned_pts.pop()

    pts = cleaned_pts

    vertices = [(x, y) for x, y, _ in pts]
    n = len(vertices)
    if vindex is not None:
        vids = [vindex.get_vid(x, y) for (x, y) in vertices]
    else:
        vids = [i + 1 for i in range(n)]
    records: List[SegmentRecord] = []

    for i in range(n):
        start_idx = i
        end_idx = (i + 1) % n
        tol_zero = 1e-6
        start_pt = vertices[start_idx]
        end_pt = vertices[end_idx]
        if _dist(start_pt, end_pt) < tol_zero:
            continue
        start_vid = vids[start_idx]
        end_vid = vids[end_idx]
        dist = _dist(start_pt, end_pt)
        az = _azimute(start_pt, end_pt)
        h_value = height_sampler(start_pt) if height_sampler is not None else 0.0
        records.append(
            SegmentRecord(
                start_vid=start_vid,
                end_vid=end_vid,
                start_point=start_pt,
                end_point=end_pt,
                start_label=_fmt_pid(start_vid),
                end_label=_fmt_pid(end_vid),
                azimute_txt=_dms_str(az),
                distancia_txt=_fmt_num(dist, 2),
                azimute_value=az,
                distancia_value=dist,
                h_value=h_value,
                h_txt=_fmt_num(h_value, 2),
            )
        )

    return vertices, vids, records


# ---------------------------------------------------------------------------
# DOCX: INSERÇÃO / PLACEHOLDERS
# ---------------------------------------------------------------------------


def _insert_paragraph_after(paragraph, text: str):
    if OxmlElement is None or Paragraph is None:
        raise RuntimeError("python-docx is required to modify the template.")
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if paragraph.style is not None:
        new_para.style = paragraph.style
    run = new_para.add_run(text)
    run.bold = False
    return new_para


def _apply_template_placeholders(doc, replacements: Dict[str, str]) -> None:
    if Document is None:
        return

    def _replace_text(paragraph) -> None:
        if paragraph.runs:
            full_text = "".join(run.text for run in paragraph.runs)
        else:
            full_text = paragraph.text
        updated_text = full_text
        for placeholder, value in replacements.items():
            if placeholder in updated_text:
                updated_text = updated_text.replace(placeholder, value)
        if updated_text == full_text:
            return
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = updated_text
            first_run.bold = False
            for run in paragraph.runs[1:]:
                run.text = ""
                run.bold = False
        else:
            paragraph.text = updated_text
            for run in paragraph.runs:
                run.bold = False

    for paragraph in doc.paragraphs:
        _replace_text(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text(paragraph)


def _format_date_label(base_label: str, today: datetime) -> str:
    month_names = [
        "janeiro",
        "fevereiro",
        "março",
        "abril",
        "maio",
        "junho",
        "julho",
        "agosto",
        "setembro",
        "outubro",
        "novembro",
        "dezembro",
    ]
    month = month_names[today.month - 1]
    formatted = base_label
    formatted = formatted.replace("yy", month)
    formatted = formatted.replace("xxx", month)
    formatted = formatted.replace("xx", f"{today.day:02d}")
    return formatted


# ---------------------------------------------------------------------------
# GERAÇÃO DOS MEMORIAIS (DOCX)
# ---------------------------------------------------------------------------


def _write_memorial_docs(
    base_dir: Path,
    poly_records: Sequence[Sequence[SegmentRecord]],
    poly_labels: Sequence[str],
    confrontantes_map: Dict[Tuple[int, int], ConfrontanteInfo],
    property_info: Sequence[PropertyInfo],
) -> None:
    if not poly_records:
        return

    if Document is None or OxmlElement is None or Paragraph is None:
        _log_error("python-docx não encontrado; os DOCs não serão gerados.")
        return
    if not DOC_TEMPLATE_PATH.exists():
        _log_error(f"Template DOCX não encontrado em '{DOC_TEMPLATE_PATH}'.")
        return

    base_dir.mkdir(parents=True, exist_ok=True)
    prefix_base = "CODEVASF"
    today = datetime.now()
    total_info = len(property_info)

    for poly_idx, (records, label) in enumerate(zip(poly_records, poly_labels), start=1):
        if not records:
            continue

        info = (
            property_info[poly_idx - 1]
            if poly_idx - 1 < total_info
            else PropertyInfo(
                label=_normalize_property_label(label, poly_idx),
                full_name=None,
                owner=None,
                document=None,
                area_ha=None,
                perimeter_m=None,
            )
        )

        prop_identifier = _normalize_property_label(info.label or label, poly_idx)
        slug = _safe_slug(prop_identifier, f"{poly_idx:02d}")
        doc_prefix = f"{prefix_base}_{slug}"

        property_dir = base_dir / doc_prefix
        property_dir.mkdir(parents=True, exist_ok=True)

        start_rec = records[0]
        start_vid = start_rec.start_vid
        start_pt = start_rec.start_point

        head = (
            f"Inicia-se a descrição deste perímetro no vértice {_fmt_pid(start_vid)}, "
            f"definido pelas coordenadas {_coord_text_colon(start_pt)};"
        )

        descr_parts: List[str] = []
        total_segments = len(records)

        for seg_idx, rec in enumerate(records):
            confrontante = _confrontante_text(confrontantes_map.get((poly_idx - 1, seg_idx)))
            use_linha_ideal = confrontante.strip().upper() == "GLEBA"

            next_vid = rec.end_vid
            next_pt = rec.end_point
            distancia = _fmt_num(rec.distancia_value, 2)
            clause = (
                f" deste, segue confrontando com {confrontante}"
                f"{' por meio de Linha ideal' if use_linha_ideal else ''} com azimute "
                f"{_short_azimute(rec.azimute_txt)} e distância {distancia} m até o vértice {_fmt_pid(next_vid)}, "
                f"definido pelas coordenadas {_coord_text_colon(next_pt)}"
            )

            if seg_idx == total_segments - 1:
                clause += ", ponto inicial da descrição, encerrando o perímetro descrito."
            else:
                clause += ";"

            descr_parts.append(clause)

        description = head + "".join(descr_parts)
        if not description.endswith("."):
            description += "."

        # Normaliza espaços em branco (remove "blank spaces" extras)
        description = _normalize_whitespace(description)

        occupant_text = info.owner.strip() if info.owner else DEFAULT_MISSING_TEXT
        document_text = info.document.strip() if info.document else DEFAULT_MISSING_TEXT
        area_value = _fmt_num(info.area_ha, 3)
        perimeter_value = _fmt_num(info.perimeter_m, 3)
        area_text = f"{area_value} ha" if area_value != "-" else area_value
        perimeter_text = f"{perimeter_value} m" if perimeter_value != "-" else perimeter_value

        doc_path = property_dir / f"{doc_prefix}.docx"

        try:
            doc = Document(str(DOC_TEMPLATE_PATH))

            desc_anchor_idx: Optional[int] = None
            date_para_idx: Optional[int] = None

            for idx_para, para in enumerate(doc.paragraphs):
                text_norm = _normalize_ascii(para.text).upper()
                if desc_anchor_idx is None and "DESCRICAO" in text_norm:
                    desc_anchor_idx = idx_para
                if date_para_idx is None and "SAO PAULO" in text_norm:
                    date_para_idx = idx_para

            if desc_anchor_idx is not None:
                desc_anchor = doc.paragraphs[desc_anchor_idx]
                next_para = doc.paragraphs[desc_anchor_idx + 1] if desc_anchor_idx + 1 < len(doc.paragraphs) else None
                if next_para is not None and not next_para.text.strip():
                    next_para.text = description
                    for run in next_para.runs:
                        run.bold = False
                else:
                    _insert_paragraph_after(desc_anchor, description)
            else:
                para = doc.add_paragraph(description)
                for run in para.runs:
                    run.bold = False

            replacements = {
                ": ID": f": {prop_identifier}",
                ":ID": f":{prop_identifier}",
                ": PRO": f": {occupant_text}",
                ":PRO": f":{occupant_text}",
                ": DOC": f": {document_text}",
                ":DOC": f":{document_text}",
                ": AREA": f": {area_text}",
                ":AREA": f":{area_text}",
                ": PER": f": {perimeter_text}",
                ":PER": f":{perimeter_text}",
            }
            _apply_template_placeholders(doc, replacements)

            if date_para_idx is not None:
                date_para = doc.paragraphs[date_para_idx]
                date_para.text = _format_date_label(date_para.text, today)
            else:
                doc.add_paragraph("\u00A0")
                doc.add_paragraph(_format_date_label("São Paulo, xx de yy de 2025", today))

            doc.save(doc_path)
        except Exception as exc:
            _log_error(f"Falha ao gerar DOCX '{doc_prefix}.docx': {exc}")


# ---------------------------------------------------------------------------
# DXFs INDIVIDUAIS POR PROPRIEDADE
# ---------------------------------------------------------------------------


def _add_polygon(ms, points: Sequence[Tuple[float, float]], layer: str, color: Optional[int] = None) -> None:
    pts = list(points)
    if len(pts) < 3:
        return
    attribs = {"layer": layer}
    if color is not None:
        attribs["color"] = color
    ms.add_lwpolyline(pts, close=True, dxfattribs=attribs)


def _write_property_dxfs(
    base_dir: Path,
    poly_vertices: Sequence[Sequence[Tuple[float, float]]],
    poly_vertices_raw: Sequence[Sequence[Tuple[float, float, float]]],
    poly_vids: Sequence[Sequence[int]],
    poly_records: Sequence[Sequence[SegmentRecord]],
    poly_headers: Sequence[Sequence[str]],
    poly_column_widths: Sequence[Sequence[float]],
    poly_labels: Sequence[str],
    property_info: Sequence[PropertyInfo],
    vertex_numbering: str,
    neighbor_span: int = 3,
) -> None:
    if not poly_vertices:
        return

    base_dir.mkdir(parents=True, exist_ok=True)

    total_polys = len(poly_vertices)
    total_info = len(property_info)

    for idx in range(total_polys):
        vertices = poly_vertices[idx]
        raw_points = poly_vertices_raw[idx] if idx < len(poly_vertices_raw) else []
        vids = poly_vids[idx]
        records = poly_records[idx]
        headers = poly_headers[idx]
        column_widths = poly_column_widths[idx]
        layer_name = poly_labels[idx]

        if not records:
            continue

        seq = idx + 1
        prop_label = _normalize_property_label(layer_name, seq)
        slug = _safe_slug(prop_label, f"{seq:02d}")
        doc_prefix = f"CODEVASF_{slug}"

        property_dir = base_dir / doc_prefix
        property_dir.mkdir(parents=True, exist_ok=True)
        dxf_path = property_dir / f"{doc_prefix}.dxf"

        params = Params()
        if DXF_TEMPLATE_PATH.exists():
            try:
                doc = ezdxf.readfile(str(DXF_TEMPLATE_PATH))
            except Exception:
                doc = ezdxf.new(setup=True)
        else:
            doc = ezdxf.new(setup=True)

        start_neighbor = max(0, idx - neighbor_span)
        end_neighbor = min(total_polys - 1, idx + neighbor_span)
        neighbor_range = range(start_neighbor, end_neighbor + 1)
        neighbor_raw_cache: Dict[int, Sequence[Tuple[float, float, float]]] = {}
        for neighbor_idx in neighbor_range:
            if neighbor_idx == idx:
                continue
            neighbor_raw = poly_vertices_raw[neighbor_idx] if neighbor_idx < len(poly_vertices_raw) else []
            neighbor_raw_cache[neighbor_idx] = neighbor_raw

        _add_basic_layers(doc)
        _register_all_blocks(doc, params)

        required_layers = (
            ("PROPRIEDADE", 7),
            ("TABELA", 7),
            ("VERTICE", 7),
            ("HACHURA", 7),
        )
        for layer_name_prop, default_color in required_layers:
            if layer_name_prop not in doc.layers:
                doc.layers.add(layer_name_prop, dxfattribs={"color": default_color})
            else:
                layer_ref = doc.layers.get(layer_name_prop)
                current_color = getattr(layer_ref.dxf, "color", 0)
                if not current_color:
                    layer_ref.dxf.color = default_color

        primary_layout = _find_layout_by_name(doc, PRIMARY_LAYOUT_NAME) or _get_preferred_layout(doc)
        secondary_layout = _find_layout_by_name(doc, SECONDARY_LAYOUT_NAME) or primary_layout
        table_layout = primary_layout

        ms = doc.modelspace()

        primary_added = False
        if raw_points:
            entity = _add_polyline_with_bulge(ms, raw_points, "PROPRIEDADE", color=7)
            primary_added = entity is not None
        if not primary_added:
            _add_polygon(ms, vertices, "PROPRIEDADE", color=7)

        try:
            hatch = ms.add_hatch(color=8, dxfattribs={"layer": "HACHURA"})
            hatch.paths.add_polyline_path([(x, y) for x, y in vertices], is_closed=True)
        except Exception:
            pass

        for neighbor_idx in neighbor_range:
            if neighbor_idx == idx:
                continue
            neighbor_raw = neighbor_raw_cache.get(neighbor_idx) or (
                poly_vertices_raw[neighbor_idx] if neighbor_idx < len(poly_vertices_raw) else []
            )
            if neighbor_raw:
                entity = _add_polyline_with_bulge(ms, neighbor_raw, "PROPRIEDADE", color=7)
                if entity is None:
                    _add_polygon(ms, poly_vertices[neighbor_idx], "PROPRIEDADE", color=7)
            elif neighbor_idx < len(poly_vertices):
                _add_polygon(ms, poly_vertices[neighbor_idx], "PROPRIEDADE", color=7)

        segmentos_dim_global: set[Tuple[int, int]] = set()

        vindex = VertexIndex(tol=0.01) if vertex_numbering == "global" else None
        today_text = datetime.now().strftime("%d/%m/%Y")

        if len(records) > TABLE_SECONDARY_THRESHOLD:
            primary_table_records = records[:TABLE_SECONDARY_THRESHOLD]
            secondary_table_records = records[TABLE_SECONDARY_THRESHOLD:]
        else:
            primary_table_records = records
            secondary_table_records = []

        if idx < total_info:
            info = property_info[idx]
        else:
            info = PropertyInfo(label=prop_label, full_name=None, owner=None, document=None, area_ha=None, perimeter_m=None)
        occupant_text = info.owner.strip() if info.owner else DEFAULT_MISSING_TEXT
        document_text = info.document.strip() if info.document else DEFAULT_MISSING_TEXT
        area_text = _fmt_num(info.area_ha, 3)
        perimeter_text = _fmt_num(info.perimeter_m, 3)
        area_display = f"{area_text} ha" if area_text != "-" else area_text
        perimeter_display = f"{perimeter_text} m" if perimeter_text != "-" else perimeter_text
        municipio_txt = info.municipio or DEFAULT_MISSING_TEXT
        uf_txt = info.uf or DEFAULT_MISSING_TEXT
        comarca_txt = info.comarca or DEFAULT_MISSING_TEXT
        cartorio_txt = info.cartorio or DEFAULT_MISSING_TEXT

        def _resolve_cns(municipio_val: str) -> str:
            norm = _normalize_ascii(municipio_val).upper()
            if "PAULO AFONSO" in norm:
                return "01.157-7"
            if "CANINDE DE SAO FRANCISCO" in norm:
                return "11.039-5"
            return DEFAULT_MISSING_TEXT

        cns_txt = _resolve_cns(municipio_txt if municipio_txt != DEFAULT_MISSING_TEXT else "")

        if vertices:
            cx, cy = _polygon_centroid(vertices)
            line_gap = max(params.altura_texto_tabela * 1.2, 1.0)
            _add_text_align(
                ms,
                f"ÁREA = {area_display}",
                cx,
                cy + line_gap / 2.0,
                params.altura_texto_tabela,
                "PROPRIEDADE",
                align="MIDDLE_CENTER",
                style=params.style_texto,
            )
            _add_text_align(
                ms,
                f"PER. = {perimeter_display}",
                cx,
                cy - line_gap / 2.0,
                params.altura_texto_tabela,
                "PROPRIEDADE",
                align="MIDDLE_CENTER",
                style=params.style_texto,
            )

        lat_txt, lon_txt, conv_txt, decl_txt, variacao_txt = (
            _compute_geo_info(vertices[0][0], vertices[0][1]) if vertices else (None, None, None, None, None)
        )
        decl_txt = decl_txt or DEFAULT_MISSING_TEXT
        variacao_txt = variacao_txt or DEFAULT_MISSING_TEXT

        bbox_prop = _poly_bbox(vertices)
        _debug(
            f"[VIEWPORT] Propriedade '{prop_label}' bbox=({bbox_prop[0]:.3f}, {bbox_prop[1]:.3f}, "
            f"{bbox_prop[2]:.3f}, {bbox_prop[3]:.3f})"
        )
        _adjust_viewports_for_property(primary_layout, bbox_prop)
        if secondary_layout is not None and secondary_layout is not primary_layout:
            _adjust_viewports_for_property(secondary_layout, bbox_prop)

        escala_txt = _get_viewport_scale(primary_layout) or DEFAULT_MISSING_TEXT
        try:
            denom = int(escala_txt.split("/")[-1])
        except Exception:
            denom = None
        if denom is not None:
            if denom <= 500:
                layout_target = _find_layout_by_name(doc, "A3")
            elif denom >= 5000:
                layout_target = _find_layout_by_name(doc, "A1")
            else:
                layout_target = None
            if layout_target is not None and layout_target is not primary_layout:
                primary_layout = layout_target
                secondary_layout = layout_target
                table_layout = layout_target
                _adjust_viewports_for_property(primary_layout, bbox_prop)
                escala_txt = _get_viewport_scale(primary_layout) or escala_txt

        table_x0, table_y0 = _table_anchor_for_layout(table_layout)
        secondary_anchor_x, secondary_anchor_y = _table_anchor_for_layout(secondary_layout)
        cartorio_primary = _cartorio_for_layout(cartorio_txt, comarca_txt, table_layout)
        cartorio_secondary = _cartorio_for_layout(cartorio_txt, comarca_txt, secondary_layout)
        _debug(
            f"Tabela posicionada em âncora primária ({table_x0:.3f}, {table_y0:.3f}) no layout "
            f"'{getattr(table_layout, 'name', 'Model') if table_layout else 'Model'}'"
        )

        _draw_memorial(
            ms,
            doc,
            params,
            list(vertices),
            list(vids),
            records,
            headers,
            vindex,
            segmentos_dim_global,
            column_widths=list(column_widths),
            vertex_layer_pt="VERTICE",
            vertex_layer_txt="VERTICE",
            table_layer="TABELA",
            use_block=True,
            property_label=prop_label,
            property_layer="PROPRIEDADE",
            table_layout=table_layout,
            x0_override=table_x0,
            y0_override=table_y0,
            table_records=primary_table_records,
        )

        if secondary_table_records:
            _draw_table_only(
                secondary_layout,
                doc,
                params,
                headers,
                list(column_widths),
                "TABELA",
                secondary_table_records,
                secondary_anchor_x,
                secondary_anchor_y,
                property_label=prop_label,
            )

        folha_primary = "01/02" if secondary_table_records else "01/01"
        positions_primary = _carimbo_positions(table_layout)
        comarca_entry_primary, cartorio_entries_primary = _cartorio_metadata_entries(
            table_layout,
            positions_primary,
            cartorio_txt,
            comarca_txt,
            municipio_txt,
            cartorio_primary,
        )
        prop_name_full = info.full_name or doc_prefix
        metadata_entries = [
            (prop_name_full, *positions_primary["prop"]),
            (occupant_text, *positions_primary["owner_main"]),
            (occupant_text, *positions_primary["owner_aux"], "MIDDLE_CENTER"),
            (document_text, *positions_primary["doc_main"]),
            (document_text, *positions_primary["doc_aux"], "MIDDLE_CENTER"),
            (today_text, *positions_primary["date"]),
            (area_display, *positions_primary["area"], "MIDDLE_CENTER"),
            (perimeter_display, *positions_primary["perimeter"], "MIDDLE_CENTER"),
            (escala_txt, *positions_primary.get("escala", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
            (lat_txt or DEFAULT_MISSING_TEXT, *positions_primary.get("lat", (0.0, 0.0)), "MIDDLE_LEFT", 2.0),
            (lon_txt or DEFAULT_MISSING_TEXT, *positions_primary.get("lon", (0.0, 0.0)), "MIDDLE_LEFT", 2.0),
            (conv_txt or DEFAULT_MISSING_TEXT, *positions_primary.get("convergencia", (0.0, 0.0)), "MIDDLE_LEFT", 2.0),
            (decl_txt, *positions_primary.get("declinacao", (0.0, 0.0)), "MIDDLE_LEFT", 2.0),
            (variacao_txt, *positions_primary.get("variacao", (0.0, 0.0)), "MIDDLE_LEFT", 2.0),
            (municipio_txt, *positions_primary.get("municipio", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
            (uf_txt, *positions_primary.get("estado", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
            comarca_entry_primary,
            *cartorio_entries_primary,
            (cns_txt, *positions_primary.get("cns", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
        ]
        _write_layout_metadata(
            doc,
            table_layout,
            params,
            metadata_entries,
            folha_primary,
            CARIMBO_LAYER,
        )
        if secondary_table_records:
            positions_secondary = _carimbo_positions(secondary_layout)
            escala_txt_secondary = _get_viewport_scale(secondary_layout) or escala_txt
            comarca_entry_secondary, cartorio_entries_secondary = _cartorio_metadata_entries(
                secondary_layout,
                positions_secondary,
                cartorio_txt,
                comarca_txt,
                municipio_txt,
                cartorio_secondary,
            )
            metadata_entries_secondary = [
                (prop_name_full, *positions_secondary["prop"]),
                (occupant_text, *positions_secondary["owner_main"]),
                (occupant_text, *positions_secondary["owner_aux"], "MIDDLE_CENTER"),
                (document_text, *positions_secondary["doc_main"]),
                (document_text, *positions_secondary["doc_aux"], "MIDDLE_CENTER"),
                (today_text, *positions_secondary["date"]),
                (area_display, *positions_secondary["area"], "MIDDLE_CENTER"),
                (perimeter_display, *positions_secondary["perimeter"], "MIDDLE_CENTER"),
                (escala_txt_secondary, *positions_secondary.get("escala", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
            (municipio_txt, *positions_secondary.get("municipio", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
            (uf_txt, *positions_secondary.get("estado", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
            comarca_entry_secondary,
            *cartorio_entries_secondary,
            (cns_txt, *positions_secondary.get("cns", (0.0, 0.0)), "MIDDLE_LEFT", 3.0),
        ]
            _write_layout_metadata(
                doc,
                secondary_layout,
                params,
                metadata_entries_secondary,
                "02/02",
                CARIMBO_LAYER,
            )

        # Remove layouts não utilizados (mantém Model + primary/secondary)
        try:
            keep_names = {
                "Model",
                getattr(primary_layout, "name", None),
                getattr(secondary_layout, "name", None),
            }
            for layout in list(doc.layouts):
                name = getattr(layout, "name", None)
                if name is None or name in keep_names:
                    continue
                try:
                    doc.layouts.delete(name)
                except Exception:
                    pass
        except Exception:
            pass

        if dxf_path.exists():
            dxf_path.unlink()
        doc.saveas(dxf_path)


# ---------------------------------------------------------------------------
# DESENHO COMPLETO DO MEMORIAL (DXF)
# ---------------------------------------------------------------------------


def _draw_memorial(
    geom_layout,
    doc: ezdxf.EzDxfDocument,
    params: Params,
    vertices: List[Tuple[float, float]],
    vids: List[int],
    records: List[SegmentRecord],
    headers: Sequence[str],
    vindex: Optional[VertexIndex],
    segmentos_dim_global: Optional[set[Tuple[int, int]]],
    column_widths: Optional[Sequence[float]] = None,
    x0_override: Optional[float] = None,
    y0_override: Optional[float] = None,
    vertex_layer_pt: str = LAYER_VERTICE_PTO,
    vertex_layer_txt: str = LAYER_VERTICE_TXT,
    table_layer: str = LAYER_TABELA,
    use_block: bool = True,
    property_label: Optional[str] = None,
    property_layer: Optional[str] = None,
    table_layout: Optional[ezdxf.layouts.BaseLayout] = None,
    table_records: Optional[Sequence[SegmentRecord]] = None,
) -> None:
    column_widths = list(column_widths) if column_widths is not None else [25.0] * len(headers)

    style_texto = params.style_texto
    ms_geom = geom_layout
    ms_table = table_layout
    layout_name = getattr(ms_table, "name", "None")
    geom_name = getattr(ms_geom, "name", "model")
    #_debug(f"Desenhando tabela em layout '{layout_name}', geometria em '{geom_name}'")

    for layer_name in (vertex_layer_pt, vertex_layer_txt, table_layer):
        if layer_name not in doc.layers:
            doc.layers.add(layer_name)

    # Desenho de vértices e rótulos Pxx
    for idx, (x, y) in enumerate(vertices):
        vid = vids[idx]
        should_draw = True if vindex is None else vindex.should_draw(vid)
        if should_draw:
            if use_block:
                ms_geom.add_blockref(BLOCO_VERTICE, (x, y), dxfattribs={"layer": vertex_layer_pt})
            else:
                ms_geom.add_circle((x, y), VERTEX_RADIUS, dxfattribs={"layer": vertex_layer_pt})

            pid_txt = _fmt_pid(vid)
            offset_x = -(VERTEX_RADIUS + abs(params.p_label_dx) + params.altura_texto_P * 1.2)
            offset_y = VERTEX_RADIUS + abs(params.p_label_dy) + params.altura_texto_P * 1.0
            label_x = x + offset_x
            label_y = y + offset_y
            _add_text_align(
                ms_geom,
                pid_txt,
                label_x,
                label_y,
                params.altura_texto_P,
                vertex_layer_txt,
                align="MIDDLE_CENTER",
                style=style_texto,
            )

            if vindex is not None:
                vindex.mark_drawn(vid)

    xs = [p[0] for p in vertices]
    ys = [p[1] for p in vertices]
    max_x, max_y = max(xs), max(ys)

    if x0_override is None or y0_override is None:
        x0 = max_x + params.tabela_offset_x
        y0 = max_y + params.tabela_offset_y
    else:
        x0, y0 = x0_override, y0_override

    table_records_seq: Sequence[SegmentRecord] = table_records if table_records is not None else records
    if ms_table is None:
        fallback_label = _fmt_pid(vids[0]) if vids else "propriedade"
        label_info = property_label or fallback_label
        #_log_error(f"Tabela nao inserida para '{label_info}' porque nenhum layout de paper space foi localizado.")
    else:
        _render_table(
            ms_table,
            params,
            headers,
            column_widths,
            table_layer,
            style_texto,
            table_records_seq,
            x0,
            y0,
        )

    segmentos_dim_local: set[Tuple[int, int]] = set()
    for rec in records:
        key = (rec.start_vid, rec.end_vid)
        if key in segmentos_dim_local:
            continue
        if segmentos_dim_global is not None:
            segmentos_dim_global.add(tuple(sorted(key)))
        segmentos_dim_local.add(key)


# ---------------------------------------------------------------------------
# PROCESSAMENTO COMPLETO DO DXF (GLOBAL + DOCS + DXFs INDIVIDUAIS)
# ---------------------------------------------------------------------------


def process_file(
    dxf_path: Path,
    layers: Optional[List[str]],
    saida_doc: Optional[Path],
    vertex_numbering: str = "global",
    generate_docs: bool = True,
    generate_dxfs: bool = True,
) -> None:
    doc = ezdxf.readfile(str(dxf_path))
    msp = doc.modelspace()

    mdt_ds = None
    height_sampler = None
    try:
        mdt_ds, height_sampler = _make_height_sampler(MDT_PATH)
    except Exception:
        mdt_ds = None
        height_sampler = None

    shp_data = _extract_polygons_from_shapefile(PROPERTIES_SHP_PATH)
    if shp_data is not None:
        shp_pts, shp_labels, shp_infos, shp_ids, road_geoms = shp_data
        polylines = [(label, pts) for label, pts in zip(shp_labels, shp_pts)]
    else:
        polylines = list(_iter_closed_polylines(msp, layers))
        road_geoms = []
        if not polylines:
            raise SystemExit("Nenhuma polyline fechada encontrada nas layers selecionadas.")

    params = Params()
    use_block = vertex_numbering == "global"
    vindex_global = VertexIndex(tol=0.01) if use_block else None

    poly_vertices: List[List[Tuple[float, float]]] = []
    poly_vertices_raw: List[List[Tuple[float, float, float]]] = []
    poly_vids: List[List[int]] = []
    poly_records: List[List[SegmentRecord]] = []
    poly_headers: List[Sequence[str]] = []
    poly_column_widths: List[List[float]] = []
    poly_labels: List[str] = []

    # Construção de registros por polilinha
    for idx, (layer, pts) in enumerate(polylines, start=1):
        pts = _rotate_start_southwest(list(pts))
        local_vindex = vindex_global if use_block else None
        vertices, vids, records = _build_segment_records(pts, local_vindex, height_sampler=height_sampler)
        poly_vertices.append(vertices)
        poly_vertices_raw.append(list(pts))
        poly_vids.append(vids)
        poly_records.append(records)

        headers = tuple(TABLE_HEADERS)
        column_widths = list(DEFAULT_COLUMN_WIDTHS)
        poly_headers.append(headers)
        poly_column_widths.append(column_widths)

        initial_label = _sanitize_label(layer) or f"{idx:02d}"
        poly_labels.append(_normalize_property_label(initial_label, idx))

    # Info básica das propriedades
    if shp_data is not None:
        poly_info = shp_infos
        id_values_numeric = shp_ids
        id_order: Optional[List[int]] = None
        if any(v is not None for v in id_values_numeric):
            id_order = sorted(
                range(len(id_values_numeric)),
                key=lambda i: (0, id_values_numeric[i]) if id_values_numeric[i] is not None else (1, i),
            )
    else:
        poly_info = [
            PropertyInfo(label=label, full_name=None, owner=None, document=None, area_ha=None, perimeter_m=None)
            for label in poly_labels
        ]
        id_order = None

    # Ordenação de polígonos e construção de confrontantes
    if poly_vertices:
        if id_order is not None:
            order = id_order
        else:
            centroids = [_polygon_centroid(vertices) for vertices in poly_vertices]
            order = _order_indices_by_chain_proximity(centroids)

        if order != list(range(len(poly_vertices))):
            poly_vertices = [poly_vertices[i] for i in order]
            poly_vertices_raw = [poly_vertices_raw[i] for i in order]
            poly_vids = [poly_vids[i] for i in order]
            poly_records = [poly_records[i] for i in order]
            poly_headers = [poly_headers[i] for i in order]
            poly_column_widths = [poly_column_widths[i] for i in order]
            poly_labels = [poly_labels[i] for i in order]
            poly_info = [poly_info[i] for i in order]

        # Seleção interativa de propriedades (opcional)
        selected_indices = _select_property_indices(poly_labels, poly_info)
        if selected_indices is not None:
            poly_vertices = [poly_vertices[i] for i in selected_indices]
            poly_vertices_raw = [poly_vertices_raw[i] for i in selected_indices]
            poly_vids = [poly_vids[i] for i in selected_indices]
            poly_records = [poly_records[i] for i in selected_indices]
            poly_headers = [poly_headers[i] for i in selected_indices]
            poly_column_widths = [poly_column_widths[i] for i in selected_indices]
            poly_labels = [poly_labels[i] for i in selected_indices]
            poly_info = [poly_info[i] for i in selected_indices]
            road_geoms = road_geoms  # unchanged

        # Anexa município/UF/comarca/cartório
        gdf_mun = _load_municipios()
        if gdf_mun is not None:
            for idx, verts in enumerate(poly_vertices):
                mun, uf, comarca, cartorio = _infer_municipio_for_poly(verts, gdf_mun)
                if idx < len(poly_info):
                    info = poly_info[idx]
                    info.municipio = mun or info.municipio
                    info.uf = uf or info.uf
                    info.comarca = comarca or info.comarca
                    info.cartorio = cartorio or info.cartorio

        confrontantes_map = _build_confrontantes_map(
            poly_records,
            poly_labels,
            poly_info,
            road_geoms,
            ROAD_INTERSECTION_TOLERANCE,
        )
    else:
        confrontantes_map = {}

    # Diretório base para saídas (DOCX/DXF)
    if saida_doc is not None:
        doc_output_dir = saida_doc if saida_doc.suffix == "" else saida_doc.parent
    else:
        doc_output_dir = dxf_path.parent

    if not generate_docs and not generate_dxfs:
        _log_error("Nenhuma saída habilitada (generate_docs=False e generate_dxfs=False); nada a fazer.")
        if mdt_ds is not None:
            try:
                mdt_ds.close()
            except Exception:
                pass
        return

    if generate_docs:
        _write_memorial_docs(doc_output_dir, poly_records, poly_labels, confrontantes_map, poly_info)
    else:
        _debug("Geração de memorial DOCX desativada; pulando.")

    if generate_dxfs:
        _write_property_dxfs(
            doc_output_dir,
            poly_vertices,
            poly_vertices_raw,
            poly_vids,
            poly_records,
            poly_headers,
            poly_column_widths,
            poly_labels,
            poly_info,
            vertex_numbering,
            neighbor_span=PROPERTY_NEIGHBOR_SPAN,
        )
    else:
        _debug("Geração de DXFs individuais desativada; pulando.")

    if mdt_ds is not None:
        try:
            mdt_ds.close()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------


def main() -> None:
    process_file(
        dxf_path=DXF_INPUT_PATH,
        layers=LAYERS_TO_PROCESS,
        saida_doc=DOC_OUTPUT_BASE,
        vertex_numbering=VERTEX_NUMBERING_MODE,
        generate_docs=GENERATE_DOCS,
        generate_dxfs=GENERATE_DXFS,
    )


if __name__ == "__main__":
    main()
