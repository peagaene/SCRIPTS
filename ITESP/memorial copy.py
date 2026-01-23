#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Gera memorial descritivo lote a lote a partir de shapefiles de LOTES e VIAS.

Premissas:
- CRS dos shapes: UTM SIRGAS 2000, fuso 22S (EPSG:31982).
- LOTES: polígonos, com campos:
    - SETOR
    - QUADRA
    - LOTE   (pode ter 009, 009A etc.)
    - LOGRADOURO
    - NUMERO
    - IMOLAD (par/ímpar) -> usado em {imolad} caso não seja possível inferir pelo número
- VIAS: linhas ou polígonos (eixo ou área), com campo:
    - N_VIA (nome da rua)

Saídas:
- Um TXT por lote, com nome: MEM_{SETOR}-{QUADRA}-{LOTE}.txt
  contendo o memorial no formato definido pelo usuário.
"""

import os
import math
from pathlib import Path

import geopandas as gpd
from shapely.geometry import Polygon, LineString, Point
from shapely.ops import unary_union, orient, substring

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
except ImportError:
    Document = None

# =========================
# CONFIGURAÇÕES DO USUÁRIO
# =========================

# Caminhos de entrada
SHAPE_LOTES = r"\\192.168.2.28\f\03 - PROCESSAMENTO\05 ITESP\05 - CODIGO\ANHUMAS\LOTES_ANHUMAS.shp"
SHAPE_VIAS = r"\\192.168.2.28\f\03 - PROCESSAMENTO\05 ITESP\05 - CODIGO\ANHUMAS\VIA.shp"
# Perímetro externo (preferencial para numeração inicial)
SHAPE_PERIMETRO = r"\\192.168.2.28\f\03 - PROCESSAMENTO\05 ITESP\05 - CODIGO\ANHUMAS\ANHUMAS.shp"
# Pontos de soleira (próximos à frente do lote)
SHAPE_SOLEIRA = r"\\192.168.2.28\f\03 - PROCESSAMENTO\05 ITESP\05 - CODIGO\ANHUMAS\SOLEIRA.shp"

# Pasta de saída para os memoriais TXT
OUT_DIR = r"\\192.168.2.28\f\03 - PROCESSAMENTO\05 ITESP\05 - CODIGO\ANHUMAS"

# Formato de saída: "docx" ou "txt"
OUTPUT_FORMAT = "docx"

# Shapefile de vértices gerado no mesmo diretório
OUT_VERTICES_SHP = os.path.join(OUT_DIR, "VERTICES.shp")
# Shapefile de linhas utilizadas na medição de distância à esquina
OUT_ESQUINA_SHP = os.path.join(OUT_DIR, "DIST_ESQUINA.shp")
# Memorial do perímetro (nome fixo)
OUT_PERIMETRO_FILE = os.path.join(OUT_DIR, "MEM_PERIMETRO.docx")
# Memorial consolidado das vias (um arquivo único)
OUT_VIAS_FILE = os.path.join(OUT_DIR, "MEM_VIAS.docx")
# Shapefile de segmentos usados (debug de confrontantes)
OUT_SEGMENTOS_SHP = os.path.join(OUT_DIR, "SEGMENTOS_DEBUG.shp")
# Shapefile de pontos médios dos segmentos (para depuração de confrontantes)
OUT_SEGMENTOS_PTO_SHP = os.path.join(OUT_DIR, "SEGMENTOS_DEBUG_PTO.shp")

# Modo de geração: "lotes", "perimetro" ou "ambos" (None pergunta no console)
MODO_GERACAO = None

# Parâmetros de interseção com vias/lotes
SEG_BUFFER = 0.05  # m: pequena folga para interseção de segmento
NEAR_VIA_MAX_DIST = 1.0  # m: fallback para achar via mais próxima se não intersectar
DEBUG_INTERSECTIONS = False  # True para imprimir avisos quando não achar via
MIN_INTER_LEN_CONFRONT = 0.0  # m: mínimo de sobreposição para considerar via como confrontante
MIN_INTER_LEN_FRENTE = 0.05  # m: mínimo de sobreposição para considerar via de acesso/esquina

# Campos da tabela de LOTES
FIELD_SETOR = "setor"
FIELD_QUADRA = "quadra"
FIELD_LOTE = "lote"
FIELD_LOGRADOURO = "logradouro"
FIELD_NUMERO = "numero"
FIELD_IMOLAD = "IMOLAD"  # par/ímpar (ex.: 'PAR', 'IMPAR') - usado como reserva

# Campo da tabela de VIAS
FIELD_N_VIA = "n_via"

# Bairro e município (valores fixos – você pode ajustar aqui)
BAIRRO_FIXO = "JARDIM PAULISTA"
MUNICIPIO_FIXO = "JUNQUEIRÓPOLIS - SP"

# EPSG do projeto (UTM SIRGAS 2000 22S)
EPSG_CRS = 31982

# Tolerância para considerar vértices iguais (em metros)
# Aumentado para 1 cm para reduzir duplicidade por ruído
VERTEX_TOL = 0.01

# =========================
# FUNÇÕES AUXILIARES
# =========================

def ensure_crs(gdf, epsg):
    """Garante que o GeoDataFrame esteja no CRS esperado (reprojeta se necessário)."""
    if gdf.crs is None:
        raise ValueError("Shapefile sem CRS definido. Defina o CRS antes de rodar o script.")
    if gdf.crs.to_epsg() != epsg:
        gdf = gdf.to_crs(epsg)
    return gdf


def key_from_coord(x, y, tol=VERTEX_TOL):
    """Cria uma chave discreta para um vértice, arredondando coordenadas."""
    return (round(x / tol) * tol, round(y / tol) * tol)


def dms_from_deg(ang):
    """Converte ângulo em graus decimais para string D°M'S\"."""
    ang = ang % 360.0
    d = int(ang)
    m_float = (ang - d) * 60.0
    m = int(m_float)
    s = (m_float - m) * 60.0
    return f"{d:03d}°{m:02d}'{s:05.2f}\"".replace(".", ",")


def azimuth_deg(x1, y1, x2, y2):
    """Calcula azimute (0-360) de (x1,y1) -> (x2,y2), a partir do norte, sentido horário."""
    dx = x2 - x1
    dy = y2 - y1
    # atan2(x, y) para referenciar Norte 0°
    ang_rad = math.atan2(dx, dy)
    ang_deg = math.degrees(ang_rad)
    if ang_deg < 0:
        ang_deg += 360.0
    return ang_deg


def distance(x1, y1, x2, y2):
    """Distância euclidiana entre dois pontos (metros)."""
    return math.hypot(x2 - x1, y2 - y1)


def fmt(num, dec=2):
    """Formata número com vírgula como separador decimal."""
    return f"{num:.{dec}f}".replace(".", ",")


def fmt(num, dec=2):
    """Formata número com vírgula como separador decimal."""
    return f"{num:.{dec}f}".replace(".", ",")


def escolher_modo():
    """Retorna ('lotes', 'perimetro', 'vias', 'ambos', 'todos')."""
    if MODO_GERACAO in ("lotes", "perimetro", "vias", "ambos", "todos"):
        return MODO_GERACAO
    try:
        choice = input("Gerar memoriais (1=lotes, 2=perímetro, 3=lotes+perímetro, 4=vias, 5=todos) [3]: ").strip()
    except Exception:
        return "ambos"
    mapa = {"1": "lotes", "2": "perimetro", "3": "ambos", "4": "vias", "5": "todos", "": "ambos"}
    return mapa.get(choice, "ambos")


def safe_filename(text):
    """Cria um nome de arquivo seguro a partir de um texto."""
    invalid = '<>:"/\\|?*'
    for ch in invalid:
        text = text.replace(ch, "_")
    text = text.replace(" ", "_")
    return text


def split_lote_code(lote_str):
    """
    Divide o campo LOTE em parte numérica + sufixo alfabético para ordenar:
    009 < 009A < 010 etc.
    """
    if lote_str is None:
        return (0, "")
    s = str(lote_str).strip()
    num_part = ""
    suf_part = ""
    for ch in s:
        if ch.isdigit():
            num_part += ch
        else:
            suf_part += ch
    num = int(num_part) if num_part else 0
    suf = suf_part.upper()
    return (num, suf)


def block_to_lines(block_list):
    """
    Converte uma lista de strings (com quebras de linha) em linhas limpas.
    Útil para alimentar o DOCX preservando conteúdo.
    """
    text = "".join(block_list)
    return [line.strip() for line in text.split("\n") if line.strip()]


def infer_imolad_from_num(numero_val, fallback=None):
    """
    Deduz o lado (PAR/IMPAR) a partir do número do imóvel.
    Se não for possível ler o número, devolve o fallback (campo IMOLAD original).
    """
    if numero_val is None:
        return fallback or ""
    num_str = str(numero_val).strip()
    digits = "".join(ch for ch in num_str if ch.isdigit())
    if not digits:
        return fallback or ""
    try:
        num_int = int(digits)
    except ValueError:
        return fallback or ""
    return "PAR" if num_int % 2 == 0 else "IMPAR"


# =========================
# LEITURA DOS DADOS
# =========================

print("Lendo shapefiles...")
gdf_lotes = gpd.read_file(SHAPE_LOTES)
gdf_vias = gpd.read_file(SHAPE_VIAS)
gdf_perimetro = None
if SHAPE_PERIMETRO and os.path.exists(SHAPE_PERIMETRO):
    gdf_perimetro = gpd.read_file(SHAPE_PERIMETRO)
gdf_soleira = None
if SHAPE_SOLEIRA and os.path.exists(SHAPE_SOLEIRA):
    gdf_soleira = gpd.read_file(SHAPE_SOLEIRA)

gdf_lotes = ensure_crs(gdf_lotes, EPSG_CRS)
gdf_vias = ensure_crs(gdf_vias, EPSG_CRS)
if gdf_perimetro is not None:
    if gdf_perimetro.crs is None:
        gdf_perimetro = gdf_perimetro.set_crs(EPSG_CRS)
    else:
        gdf_perimetro = ensure_crs(gdf_perimetro, EPSG_CRS)
if gdf_soleira is not None:
    if gdf_soleira.crs is None:
        gdf_soleira = gdf_soleira.set_crs(EPSG_CRS)
    else:
        gdf_soleira = ensure_crs(gdf_soleira, EPSG_CRS)

HAS_IMOLAD_FIELD = FIELD_IMOLAD in gdf_lotes.columns
gdf_lotes["_orig_index"] = gdf_lotes.index

# Ordenação de lotes (códigos alfanuméricos)
gdf_lotes["_lote_order"] = gdf_lotes[FIELD_LOTE].apply(split_lote_code)
gdf_lotes_sorted = gdf_lotes.sort_values(
    by=[FIELD_SETOR, FIELD_QUADRA, "_lote_order"]
).reset_index(drop=True)

# Garante que vias sejam LINESTRING (se forem polígonos, usa o contorno)
if gdf_vias.geometry.iloc[0].geom_type.startswith("Poly"):
    gdf_vias["geometry"] = gdf_vias.geometry.boundary

# =========================
# TABELA GLOBAL DE VÉRTICES
# =========================

print("Construindo tabela global de vértices (perímetro + internos)...")

# 1) Perímetro do parcelamento (preferir shape externo; senão união dos lotes)
if gdf_perimetro is not None and not gdf_perimetro.empty:
    perimeter_geom = unary_union(gdf_perimetro.geometry)
    print("Perímetro externo carregado de SHAPE_PERIMETRO.")
else:
    perimeter_geom = unary_union(gdf_lotes.geometry)
    print("SHAPE_PERIMETRO não encontrado/vazio; usando união dos lotes.")

perimeter_keys = set()
perimeter_vertices = []

def ring_to_clockwise(coords):
    # Remove último repetido
    if len(coords) > 1 and coords[0] == coords[-1]:
        coords = coords[:-1]
    # Checa orientação pelo sinal da área (shoelace): >0 ccw, <0 cw
    area2 = 0.0
    n = len(coords)
    for i in range(n):
        x1, y1 = coords[i][0], coords[i][1]
        x2, y2 = coords[(i + 1) % n][0], coords[(i + 1) % n][1]
        area2 += (x1 * y2 - x2 * y1)
    if area2 > 0:  # contra-horário, inverter para horário
        coords = list(reversed(coords))
    # Rotaciona para começar no ponto mais ao norte (y max; empate por x min)
    max_idx = max(range(len(coords)), key=lambda i: (coords[i][1], -coords[i][0]))
    coords = coords[max_idx:] + coords[:max_idx]
    return coords

def collect_perimeter_coords(geom):
    if geom.is_empty:
        return
    if geom.geom_type == "Polygon":
        rings = [geom.exterior]
    elif geom.geom_type == "MultiPolygon":
        rings = [p.exterior for p in geom.geoms]
    else:
        return
    for ring in rings:
        coords = ring_to_clockwise(list(ring.coords))
        for coord in coords:
            x, y = coord[0], coord[1]
            key = key_from_coord(x, y)
            if key not in perimeter_keys:
                perimeter_keys.add(key)
                perimeter_vertices.append((key, x, y))

collect_perimeter_coords(perimeter_geom)

# 2) Todos os vértices dos lotes (internos + perímetro)
all_vertex_dict = {}  # key -> {"x":, "y":, "is_perimeter": bool, "code": None}
vertex_lot_rank = {}

# Marca primeiro os de perímetro
for key, x, y in perimeter_vertices:
    all_vertex_dict[key] = {"x": x, "y": y, "is_perimeter": True, "code": None}

# Agora varre todos os lotes (ordenados) e adiciona os demais
for lot_idx, geom in enumerate(gdf_lotes_sorted.geometry):
    if geom.is_empty:
        continue
    if isinstance(geom, Polygon):
        polys = [geom]
    else:
        polys = [p for p in geom.geoms if isinstance(p, Polygon)]
    for poly in polys:
        coords = list(poly.exterior.coords)
        if len(coords) > 1 and coords[0] == coords[-1]:
            coords = coords[:-1]
        for coord in coords:
            x, y = coord[0], coord[1]
            key = key_from_coord(x, y)
            if key not in all_vertex_dict:
                all_vertex_dict[key] = {"x": x, "y": y, "is_perimeter": False, "code": None}
            vertex_lot_rank[key] = min(vertex_lot_rank.get(key, lot_idx), lot_idx)

# 3) Numeração: perímetro primeiro (mais ao norte), depois internos
# Ordena perímetro conforme a sequência do anel (clockwise a partir do norte)
perim_sorted = [k for k, _, _ in perimeter_vertices]
# Internos: ordena por lote (ordem de gdf_lotes_sorted), depois norte->sul, x crescente
intern_sorted = sorted(
    [k for k, v in all_vertex_dict.items() if not v["is_perimeter"]],
    key=lambda k: (
        vertex_lot_rank.get(k, float("inf")),
        -all_vertex_dict[k]["y"],
        all_vertex_dict[k]["x"],
    ),
)

current_id = 1
for key in perim_sorted + intern_sorted:
    all_vertex_dict[key]["code"] = f"P-{current_id:03d}"
    current_id += 1

# Adiciona vértices das vias (continua numeração)
for geom in gdf_vias.geometry:
    if geom.is_empty:
        continue
    if geom.geom_type == "LineString":
        coords = list(geom.coords)
    elif geom.geom_type == "MultiLineString":
        coords = []
        for ln in geom.geoms:
            coords.extend(list(ln.coords))
    else:
        continue
    for coord in coords:
        x, y = coord[0], coord[1]
        key = key_from_coord(x, y)
        if key not in all_vertex_dict:
            all_vertex_dict[key] = {"x": x, "y": y, "is_perimeter": False, "code": f"P-{current_id:03d}"}
            current_id += 1

print(f"Total de vértices numerados: {len(all_vertex_dict)}")

# Índice espacial para vias e lotes (para confrontantes)
vias_sindex = gdf_vias.sindex
lotes_sindex = gdf_lotes.sindex

# =========================
# FUNÇÕES PARA CONFRONTANTES
# =========================

def get_via_name_for_segment(seg: LineString, exclude_vias=None):
    """Retorna o nome da via que mais intercepta o segmento (se houver), ignorando vias em exclude_vias."""
    exclude_vias = set(exclude_vias or [])
    buf = seg.buffer(SEG_BUFFER)
    possible = list(vias_sindex.intersection(buf.bounds))
    # prioriza via mais próxima do ponto médio, exigindo sobreposição mínima; desempata por maior interseção
    mid_pt = seg.interpolate(0.5, normalized=True)
    best = (None, float("inf"), 0.0)  # (via, dist_mid, inter_len)
    for idx in possible:
        via_geom = gdf_vias.geometry.iloc[idx]
        via_name = gdf_vias.iloc[idx][FIELD_N_VIA]
        if via_name in exclude_vias:
            continue
        if not via_geom.intersects(buf):
            continue
        inter_len = seg.intersection(via_geom).length
        if inter_len < MIN_INTER_LEN_CONFRONT:
            continue
        dist_mid = mid_pt.distance(via_geom)
        if (dist_mid < best[1]) or (
            math.isclose(dist_mid, best[1], abs_tol=1e-6) and inter_len > best[2]
        ):
            best = (via_name, dist_mid, inter_len)
    if best[0]:
        return best[0]
    # Fallback: via mais próxima dentro de NEAR_VIA_MAX_DIST
    nearest_name = None
    nearest_dist = None
    for idx in possible:
        via_geom = gdf_vias.geometry.iloc[idx]
        via_name = gdf_vias.iloc[idx][FIELD_N_VIA]
        if via_name in exclude_vias:
            continue
        d = seg.distance(via_geom)
        if (nearest_dist is None or d < nearest_dist) and d <= NEAR_VIA_MAX_DIST:
            nearest_dist = d
            nearest_name = via_name
    if nearest_name and DEBUG_INTERSECTIONS:
        print(f"[DEBUG] Via aproximada usada (dist {nearest_dist:.2f} m): {nearest_name}")
    return nearest_name


def get_neighbor_lote_for_segment(seg: LineString, this_orig_index):
    """Retorna (setor, quadra, lote, inter_len) do lote vizinho mais aderente."""
    buf = seg.buffer(0.05)
    possible = list(lotes_sindex.intersection(buf.bounds))
    best = (None, 0.0)
    for idx in possible:
        if idx == this_orig_index:
            continue
        lote_geom = gdf_lotes.geometry.iloc[idx]
        if not lote_geom.intersects(buf):
            continue
        inter_len = seg.intersection(lote_geom).length
        if inter_len > best[1]:
            row = gdf_lotes.iloc[idx]
            best = (
                (row[FIELD_SETOR], row[FIELD_QUADRA], row[FIELD_LOTE]),
                inter_len,
            )
    return best[0], best[1]


def build_conf_text(seg: LineString, this_orig_index, via_atual=None):
    viz, viz_len = get_neighbor_lote_for_segment(seg, this_orig_index)
    via_name = get_via_name_for_segment(seg, exclude_vias={via_atual} if via_atual else None)
    via_len = 0.0
    if via_name:
        via_geom_sel = gdf_vias[gdf_vias[FIELD_N_VIA] == via_name]
        if not via_geom_sel.empty:
            via_len = seg.intersection(via_geom_sel.geometry.union_all()).length

    if viz and viz_len >= via_len:
        setor, quadra, lote = viz
        return f"confrontando com o Lote {lote} da Quadra {quadra}"
    if via_name:
        if via_atual:
            return f"intersectando com a {via_name}"
        return f"confrontando com a {via_name}"
    return "confrontante não identificado"


def get_vias_de_acesso(geom: Polygon):
    """
    Retorna as vias que intersectam o lote, ordenadas por comprimento
    de interseção (maior primeiro). Se nenhuma intersectar acima do
    mínimo, usa as mais próximas dentro de NEAR_VIA_MAX_DIST.
    """
    results = []
    results_near = []
    for idx, via in gdf_vias.iterrows():
        inter = geom.intersection(via.geometry)
        if not inter.is_empty:
            length = inter.length
            if length >= MIN_INTER_LEN_FRENTE:
                results.append((via[FIELD_N_VIA], length))
        # fallback: via mais próxima dentro de NEAR_VIA_MAX_DIST
        d = geom.distance(via.geometry)
        if d <= NEAR_VIA_MAX_DIST:
            results_near.append((via[FIELD_N_VIA], 0.0, d))

    if results:
        results.sort(key=lambda x: -x[1])
        return [v for v, _ in results]

    # sem interseção acima do mínimo: usa proximidade
    if results_near:
        results_near.sort(key=lambda x: x[2])  # menor distância
        return [v for v, _, _ in results_near]

    return []


def is_lote_esquina(geom: Polygon):
    """
    True se o lote intersecta (acima do mínimo) duas vias com nomes diferentes.
    """
    vias = get_vias_de_acesso(geom)
    return len(set(vias[:2])) >= 2


def distancia_da_esquina(geom: Polygon, via_nome):
    """
    Calcula distância aproximada da frente até a esquina:
    - pega endpoints da interseção lote x via (cantos da testada)
    - pega pontos da frente (interseção via x fronteira do lote)
    - mede menor distância linear ao longo da via entre um canto e um ponto de frente
    Retorna (distância, linha usada)
    """
    via_rows = gdf_vias[gdf_vias[FIELD_N_VIA] == via_nome]
    if via_rows.empty:
        return None, None

    via_geom = via_rows.geometry.union_all()
    inter = geom.intersection(via_geom)
    if inter.is_empty:
        return None, None

    # Candidatos a esquina: endpoints dos trechos de interseção
    corner_points = []
    if inter.geom_type == "LineString":
        corner_points.extend([Point(inter.coords[0]), Point(inter.coords[-1])])
    elif inter.geom_type == "MultiLineString":
        for line in inter.geoms:
            corner_points.extend([Point(line.coords[0]), Point(line.coords[-1])])

    # Pontos da frente: interseção via x fronteira do lote
    front_geom = via_geom.intersection(geom.boundary)
    front_points = []
    if front_geom.is_empty:
        return None, None
    if front_geom.geom_type == "Point":
        front_points.append(front_geom)
    elif front_geom.geom_type == "MultiPoint":
        front_points.extend(list(front_geom.geoms))
    elif front_geom.geom_type in ("LineString", "MultiLineString"):
        geoms = [front_geom] if front_geom.geom_type == "LineString" else list(front_geom.geoms)
        for ln in geoms:
            front_points.append(Point(ln.coords[0]))
            front_points.append(Point(ln.coords[-1]))

    if not corner_points or not front_points:
        return None, None

    def proj_dist(pt):
        return via_geom.project(pt)

    best = (None, None)
    for corner in corner_points:
        c_proj = proj_dist(corner)
        for front in front_points:
            f_proj = proj_dist(front)
            dist_lin = abs(f_proj - c_proj)
            try:
                seg_line = substring(via_geom, min(f_proj, c_proj), max(f_proj, c_proj))
            except Exception:
                seg_line = LineString([(corner.x, corner.y), (front.x, front.y)])
            if (best[0] is None) or (dist_lin < best[0]):
                best = (dist_lin, seg_line)

    return best


def rotate_coords_starting_on_via(coords, via_name):
    """
    Rotaciona a lista de coordenadas (já em sentido horário, sem ponto repetido)
    para começar no vértice mais ao norte do segmento que intercepta a via de acesso.
    """
    if not coords or via_name is None:
        return coords
    via_rows = gdf_vias[gdf_vias[FIELD_N_VIA] == via_name]
    if via_rows.empty:
        return coords
    via_geom = via_rows.geometry.union_all()
    n = len(coords)
    best_idx = None
    best_y = None
    best_x = None
    for i in range(n):
        p1 = Point(coords[i])
        p2 = Point(coords[(i + 1) % n])
        seg = LineString([p1, p2])
        if not seg.intersects(via_geom):
            continue
        for idx in (i, (i + 1) % n):
            x, y = coords[idx][0], coords[idx][1]
            if (best_idx is None) or (y > best_y) or (y == best_y and x < best_x):
                best_idx = idx
                best_y = y
                best_x = x
    if best_idx is None:
        return coords
    return coords[best_idx:] + coords[:best_idx]


def rotate_coords_starting_on_front(coords, geom, via_name):
    """
    Tenta usar soleira para localizar a frente; se não houver, cai para a via.
    """
    if not coords:
        return coords
    best_idx = None
    # Usa ponto de soleira mais próximo do lote (se disponível)
    if gdf_soleira is not None and not gdf_soleira.empty:
        try:
            dists = gdf_soleira.geometry.distance(geom)
            idx = dists.idxmin()
            pt = gdf_soleira.geometry.loc[idx]
            min_seg_dist = None
            n = len(coords)
            for i in range(n):
                p1 = Point(coords[i])
                p2 = Point(coords[(i + 1) % n])
                seg = LineString([p1, p2])
                d = seg.distance(pt)
                if (min_seg_dist is None) or (d < min_seg_dist):
                    min_seg_dist = d
                    best_idx = i if coords[i][1] >= coords[(i + 1) % n][1] else (i + 1) % n
        except Exception:
            best_idx = None
    if best_idx is not None:
        return coords[best_idx:] + coords[:best_idx]
    # Fallback: usa via
    return rotate_coords_starting_on_via(coords, via_name)


# =========================
# GERAÇÃO DOS MEMORIAIS
# =========================

Path(OUT_DIR).mkdir(parents=True, exist_ok=True)

modo = escolher_modo()
RUN_LOTES = modo in ("lotes", "ambos", "todos")
RUN_PERIMETRO = modo in ("perimetro", "ambos", "todos")
RUN_VIAS = modo in ("vias", "ambos", "todos") if isinstance(modo, str) else True
segment_records = []
segment_mid_records = []

# Exporta shapefile com todos os vértices numerados
vertices_records = []
for data in all_vertex_dict.values():
    vertices_records.append(
        {
            "code": data["code"],
            "is_perimeter": data["is_perimeter"],
            "x": data["x"],
            "y": data["y"],
            "geometry": Point(data["x"], data["y"]),
        }
    )
gdf_vertices = gpd.GeoDataFrame(vertices_records, crs=f"EPSG:{EPSG_CRS}")
gdf_vertices.to_file(OUT_VERTICES_SHP)
print(f"Shapefile de vértices gerado: {OUT_VERTICES_SHP}")

# Ordena os lotes por SETOR, QUADRA, LOTE (009 antes de 009A)
gdf_lotes["_lote_order"] = gdf_lotes[FIELD_LOTE].apply(split_lote_code)
gdf_lotes_sorted = gdf_lotes.sort_values(
    by=[FIELD_SETOR, FIELD_QUADRA, "_lote_order"]
).reset_index(drop=True)

if RUN_LOTES:
    print("Gerando memoriais lote a lote...")
    
    esquina_records = []
    
    for idx, row in gdf_lotes_sorted.iterrows():
        geom = row.geometry
        if geom.is_empty:
            continue
    
        # Garante polígono e sentido horário
        if not isinstance(geom, Polygon):
            # se for MultiPolygon, pega o maior
            polys = [p for p in geom.geoms if isinstance(p, Polygon)]
            polys.sort(key=lambda p: p.area, reverse=True)
            geom = polys[0]
    
        poly_cw = orient(geom, sign=-1.0)  # força sentido horário
        coords = list(poly_cw.exterior.coords)
        if len(coords) > 1 and coords[0] == coords[-1]:
            coords = coords[:-1]
        setor = row[FIELD_SETOR]
        quadra = row[FIELD_QUADRA]
        lote = row[FIELD_LOTE]
        logradouro = row[FIELD_LOGRADOURO]
        numero = row[FIELD_NUMERO]
        # Calcula imolad pelo número do imóvel; se falhar, usa o valor original do campo (se existir)
        imolad_fallback = row[FIELD_IMOLAD] if HAS_IMOLAD_FIELD else None
        imolad = infer_imolad_from_num(numero, fallback=imolad_fallback)

        # Rotaciona para iniciar na frente (via de acesso): vértice mais ao norte do segmento que intercepta a via
        vias_lote = get_vias_de_acesso(geom)
        via_acesso = vias_lote[0] if vias_lote else logradouro
        coords = rotate_coords_starting_on_via(coords, via_acesso)

        # Cabeçalho
        # Reavalia frente do lote para alinhar com soleira/via e via de acesso
        vias_lote = get_vias_de_acesso(geom)
        via_preferida = vias_lote[0] if vias_lote else logradouro
        coords = rotate_coords_starting_on_front(coords, geom, via_preferida)
        seg0_line = LineString([coords[0], coords[1]]) if len(coords) >= 2 else None
        via_acesso = get_via_name_for_segment(seg0_line) if seg0_line else None
        if not via_acesso:
            via_acesso = via_preferida
    
        imo_id = f"{setor}-{quadra}-{lote}"
        imoend = f"{logradouro}, nº {numero}"
        imoare = geom.area  # m²
    
        # Texto de esquinas / distância
        imoesqext = ""
        if is_lote_esquina(geom) and len(vias_lote) >= 2:
            imoesqext = (
                f", tratando-se de lote de esquina, com testada principal voltada para {vias_lote[1]} "
                f"e testada secundária voltada para {vias_lote[0]}"
            )
        elif vias_lote:
            # distância fixa 0,00 m para todos os lotes (texto mantido)
            dist_esq = 0.0
            imoesqext = f", a {fmt(dist_esq,2)} m da esquina mais próxima com a {vias_lote[0]}"
    
        # Vértice inicial: primeiro da lista (já está em sentido horário)
        x0, y0 = coords[0][0], coords[0][1]
        key0 = key_from_coord(x0, y0)
        vert_code0 = all_vertex_dict[key0]["code"]
    
        # Corpo da descrição: segmentos
        segmentos = []
        n = len(coords)
        for i in range(n):
            x1, y1 = coords[i][0], coords[i][1]
            x2, y2 = coords[(i + 1) % n][0], coords[(i + 1) % n][1]
            key1 = key_from_coord(x1, y1)
            key2 = key_from_coord(x2, y2)
            code1 = all_vertex_dict[key1]["code"]
            code2 = all_vertex_dict[key2]["code"]

            seg_line = LineString([(x1, y1), (x2, y2)])
            via_used = get_via_name_for_segment(seg_line)
            conf_text = build_conf_text(seg_line, row["_orig_index"])

            dist = distance(x1, y1, x2, y2)
            az = azimuth_deg(x1, y1, x2, y2)
            az_str = dms_from_deg(az)

            segmentos.append(
                {
                    "from_code": code1,
                    "to_code": code2,
                    "from_x": x1,
                    "from_y": y1,
                    "to_x": x2,
                    "to_y": y2,
                    "dist": dist,
                    "az": az_str,
                    "conf": conf_text,
                }
            )
            segment_records.append(
                {
                    "tipo": "lote",
                    "setor": setor,
                    "quadra": quadra,
                    "lote": lote,
                    "from": code1,
                    "to": code2,
                    "via": via_used,
                    "conf": conf_text,
                    "geometry": seg_line,
                }
            )
            mid_pt = seg_line.interpolate(0.5, normalized=True)
            segment_mid_records.append(
                {
                    "tipo": "lote",
                    "setor": setor,
                    "quadra": quadra,
                    "lote": lote,
                    "from": code1,
                    "to": code2,
                    "via": via_used,
                    "conf": conf_text,
                    "geometry": Point(mid_pt.x, mid_pt.y),
                }
            )
    
        # Montagem do texto conforme modelo
        header = []
        # Cabeçalho institucional centralizado
        instit_block = [
            "GOVERNO DO ESTADO DE SÃO PAULO",
            "SECRETARIA DA JUSTIÇA E CIDADANIA",
            "FUNDAÇÃO INSTITUTO DE TERRAS DO ESTADO DE SÃO PAULO",
            '"JOSÉ GOMES DA SILVA" - ITESP',
            "",
            "MEMORIAL DESCRITIVO",
            "",
        ]
        for line in instit_block:
            header.append(f"{line.center(80)}\n")
        header.append(f"SETOR: {setor} - QUADRA: {quadra} - LOTE: {lote}\n")
        header.append(f"LOGRADOURO: {imoend}\n")
        header.append(f"BAIRRO: {BAIRRO_FIXO}\n")
        header.append(f"MUNICÍPIO: {MUNICIPIO_FIXO}\n")
        header.append(f"ÁREA: {fmt(imoare,2)} m²\n\n")
    
        header.append(
            f"O lote que ora se descreve, para quem de frente olha para o imóvel, "
            f"encontra-se no lado {imolad} da {via_acesso}{imoesqext}.\n\n"
        )
    
        descricao = []
        seg0 = segmentos[0]
        descricao.append(
            "DESCRIÇÃO\n\n"
            f"Inicia-se a descrição deste perímetro no vértice {seg0['from_code']}, "
            f"de coordenadas N {fmt(seg0['from_y'],3)} m e E {fmt(seg0['from_x'],3)} m; "
        )
    
        # Segmentos intermediários
        for seg in segmentos[:-1]:
            descricao.append(
                f"deste, segue {seg['conf']} com os seguintes azimutes e distâncias: "
                f"{seg['az']} e {fmt(seg['dist'],2)} m até o vértice {seg['to_code']}, "
                f"de coordenadas N {fmt(seg['to_y'],3)} m e E {fmt(seg['to_x'],3)} m; "
            )
    
        # Último segmento (fecha no vértice inicial)
        seg_last = segmentos[-1]
        descricao.append(
            f"por fim, segue {seg_last['conf']} com os seguintes azimutes e distâncias: "
            f"{seg_last['az']} e {fmt(seg_last['dist'],2)} m até o vértice {seg_last['to_code']}, "
            "ponto inicial da descrição deste perímetro. "
        )
    
        descricao.append(
            "Todas as coordenadas aqui descritas estão georreferenciadas ao Sistema "
            "Geodésico Brasileiro, tendo como datum SIRGAS 2000 e sistema de projeção UTM, fuso 22S.\n"
        )
        descricao.append(
            "Todos os azimutes e distâncias, área e perímetro foram calculados no sistema "
            "local de coordenadas com origem do plano definido pela média das coordenadas "
            "(SGL – Sistema Geodésico Local).\n"
        )
    
        texto_final = "".join(header) + "".join(descricao)
    
        output_fmt = OUTPUT_FORMAT.lower()
        if output_fmt == "docx":
            if Document is None:
                raise ImportError("python-docx não instalado. Instale para gerar DOCX.")
            doc = Document()
            # Margens: topo/esquerda 3 cm, base/direita 2 cm (aplicado a todas as seções).
            # python-docx não oferece alternar margens em páginas pares/ímpares; use impressão duplex com "espelhar margens" se precisar.
            for section in doc.sections:
                section.top_margin = Cm(3)
                section.left_margin = Cm(3)
                section.bottom_margin = Cm(2)
                section.right_margin = Cm(2)
    
            header_lines = block_to_lines(header)
            after_memorial = False
            for line in header_lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                if line.upper().startswith("MEMORIAL DESCRITIVO"):
                    run.bold = True
                if line.startswith("O lote que ora se descreve"):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif after_memorial:
                    p.alignment = 0  # left
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if line.upper().startswith("MEMORIAL DESCRITIVO"):
                        after_memorial = True
    
            for paragraph in "".join(descricao).split("\n\n"):
                text = paragraph.replace("\n", " ").strip()
                if text:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    if text.upper().startswith("DESCRIÇÃO"):
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
            fname = safe_filename(f"MEM_{setor}-{quadra}-{lote}.docx")
            fpath = os.path.join(OUT_DIR, fname)
            doc.save(fpath)
        else:
            fname = safe_filename(f"MEM_{setor}-{quadra}-{lote}.txt")
            fpath = os.path.join(OUT_DIR, fname)
            with open(fpath, "w", encoding="utf-8") as f:
                f.write(texto_final)
    
        print(f"Memorial gerado: {fpath}")
    
    # =========================
# MEMORIAL DO PERÍMETRO
# =========================
if perimeter_geom and not perimeter_geom.is_empty:
    print("Gerando memorial do perímetro...")
    per_poly = perimeter_geom
    if per_poly.geom_type != "Polygon":
        if per_poly.geom_type == "MultiPolygon":
            per_poly = max(per_poly.geoms, key=lambda g: g.area)
        else:
            per_poly = None
    if per_poly:
        coords = [(c[0], c[1]) for c in orient(per_poly, sign=-1.0).exterior.coords]
        if len(coords) > 1 and coords[0] == coords[-1]:
            coords = coords[:-1]
        # inicia no ponto mais ao norte
        start_idx = max(range(len(coords)), key=lambda i: (coords[i][1], -coords[i][0]))
        coords = coords[start_idx:] + coords[:start_idx]
        setor_per = ""
        if FIELD_SETOR in gdf_lotes.columns:
            try:
                setor_per = str(gdf_lotes[FIELD_SETOR].iloc[0])
            except Exception:
                setor_per = ""

        segmentos = []
        n = len(coords)
        for i in range(n):
            x1, y1 = coords[i][0], coords[i][1]
            x2, y2 = coords[(i + 1) % n][0], coords[(i + 1) % n][1]
            seg_line = LineString([(x1, y1), (x2, y2)])
            via_name = get_via_name_for_segment(seg_line) or "via não identificada"
            dist = distance(x1, y1, x2, y2)
            az = azimuth_deg(x1, y1, x2, y2)
            segmentos.append(
                {
                    "from_code": f"P-{i+1:03d}",
                    "to_code": f"P-{(i+2) if (i+1)<n else 1:03d}",
                    "from_x": x1,
                    "from_y": y1,
                    "to_x": x2,
                    "to_y": y2,
                    "dist": dist,
                    "az": dms_from_deg(az),
                    "conf": f"confrontando com a {via_name}",
                }
            )
            segment_records.append(
                {
                    "tipo": "perimetro",
                    "setor": setor_per,
                    "from": segmentos[-1]["from_code"],
                    "to": segmentos[-1]["to_code"],
                    "via": via_name,
                    "conf": f"confrontando com a {via_name}",
                    "geometry": seg_line,
                }
            )
            mid_pt = seg_line.interpolate(0.5, normalized=True)
            segment_mid_records.append(
                {
                    "tipo": "perimetro",
                    "setor": setor_per,
                    "from": segmentos[-1]["from_code"],
                    "to": segmentos[-1]["to_code"],
                    "via": via_name,
                    "conf": f"confrontando com a {via_name}",
                    "geometry": Point(mid_pt.x, mid_pt.y),
                }
            )

        header = []
        instit_block = [
            "GOVERNO DO ESTADO DE SÃO PAULO",
            "SECRETARIA DA JUSTIÇA E CIDADANIA",
            "FUNDAÇÃO INSTITUTO DE TERRAS DO ESTADO DE SÃO PAULO",
            '"JOSÉ GOMES DA SILVA" - ITESP',
            "",
            "MEMORIAL DESCRITIVO",
            "",
        ]
        for line in instit_block:
            header.append(f"{line.center(80)}\n")
        header.append(f"PERÍMETRO DO NÚCLEO {setor_per}\n")
        header.append(f"BAIRRO: {BAIRRO_FIXO}\n")
        header.append(f"MUNICÍPIO: {MUNICIPIO_FIXO}\n")
        header.append(f"ÁREA: {fmt(per_poly.area,2)} m²\n\n")

        seg0 = segmentos[0]
        descricao = []
        descricao.append(
            "DESCRIÇÃO\n\n"
            f"Inicia-se a descrição deste perímetro no vértice {seg0['from_code']}, "
            f"de coordenadas N {fmt(seg0['from_y'],3)} m e E {fmt(seg0['from_x'],3)} m; "
        )
        for seg in segmentos[:-1]:
            descricao.append(
                f"deste, segue {seg['conf']} com os seguintes azimutes e distâncias: "
                f"{seg['az']} e {fmt(seg['dist'],2)} m até o vértice {seg['to_code']}, "
                f"de coordenadas N {fmt(seg['to_y'],3)} m e E {fmt(seg['to_x'],3)} m; "
            )
        seg_last = segmentos[-1]
        descricao.append(
            f"por fim, segue {seg_last['conf']} com os seguintes azimutes e distâncias: "
            f"{seg_last['az']} e {fmt(seg_last['dist'],2)} m até o vértice {seg_last['to_code']}, "
            "ponto inicial da descrição deste perímetro. "
        )
        descricao.append(
            "Todas as coordenadas aqui descritas estão georreferenciadas ao Sistema "
            "Geodésico Brasileiro, tendo como datum SIRGAS 2000 e sistema de projeção UTM, fuso 22S.\n"
        )
        descricao.append(
            "Todos os azimutes e distâncias, área e perímetro foram calculados no sistema "
            "local de coordenadas com origem do plano definido pela média das coordenadas "
            "(SGL – Sistema Geodésico Local).\n"
        )

        texto_final = "".join(header) + "".join(descricao)

        if OUTPUT_FORMAT.lower() == "docx":
            if Document is None:
                raise ImportError("python-docx não instalado. Instale para gerar DOCX.")
            doc = Document()
            # margens
            for section in doc.sections:
                section.top_margin = Cm(3)
                section.left_margin = Cm(3)
                section.bottom_margin = Cm(2)
                section.right_margin = Cm(2)
            header_lines = block_to_lines(header)
            after_memorial = False
            for line in header_lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                if line.upper().startswith("MEMORIAL DESCRITIVO"):
                    run.bold = True
                if after_memorial:
                    p.alignment = 0
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if line.upper().startswith("MEMORIAL DESCRITIVO"):
                        after_memorial = True
            for paragraph in "".join(descricao).split("\n\n"):
                text = paragraph.replace("\n", " ").strip()
                if text:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    if text.upper().startswith("DESCRIÇÃO"):
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.save(OUT_PERIMETRO_FILE)
        else:
            with open(OUT_PERIMETRO_FILE.replace(".docx", ".txt"), "w", encoding="utf-8") as f:
                f.write(texto_final)
        print(f"Memorial do perímetro gerado: {OUT_PERIMETRO_FILE}")


if RUN_VIAS:
    print("Gerando memoriais das vias...")
    output_fmt = OUTPUT_FORMAT.lower()
    doc_vias = None
    vias_txt_parts = []
    if output_fmt == "docx":
        if Document is None:
            raise ImportError("python-docx não instalado. Instale para gerar DOCX.")
        doc_vias = Document()
        for section in doc_vias.sections:
            section.top_margin = Cm(3)
            section.left_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.right_margin = Cm(2)
    via_count = 0

    for via_nome, via_group in gdf_vias.groupby(FIELD_N_VIA):
        via_parts = []
        for geom in via_group.geometry:
            if geom.is_empty:
                continue
            lines = [geom] if geom.geom_type == "LineString" else list(geom.geoms)
            for ln in lines:
                coords = list(ln.coords)
                if len(coords) < 2:
                    continue
                segmentos_via = []
                for i in range(len(coords) - 1):
                    x1, y1 = coords[i][0], coords[i][1]
                    x2, y2 = coords[i + 1][0], coords[i + 1][1]
                    key1 = key_from_coord(x1, y1)
                    key2 = key_from_coord(x2, y2)
                    code1 = all_vertex_dict[key1]["code"] if key1 in all_vertex_dict else f"P-{current_id:03d}"
                    if key1 not in all_vertex_dict:
                        all_vertex_dict[key1] = {"x": x1, "y": y1, "is_perimeter": False, "code": code1}
                        current_id += 1
                    code2 = all_vertex_dict[key2]["code"] if key2 in all_vertex_dict else f"P-{current_id:03d}"
                    if key2 not in all_vertex_dict:
                        all_vertex_dict[key2] = {"x": x2, "y": y2, "is_perimeter": False, "code": code2}
                        current_id += 1
                    seg_line = LineString([(x1, y1), (x2, y2)])
                    conf_text = build_conf_text(seg_line, this_orig_index=-1, via_atual=via_nome)
                    dist = distance(x1, y1, x2, y2)
                    az = azimuth_deg(x1, y1, x2, y2)
                    segmentos_via.append(
                        {
                            "from_code": code1,
                            "to_code": code2,
                            "from_x": x1,
                            "from_y": y1,
                            "to_x": x2,
                            "to_y": y2,
                            "dist": dist,
                            "az": dms_from_deg(az),
                            "conf": conf_text,
                        }
                    )
                if segmentos_via:
                    via_parts.append(segmentos_via)

        if not via_parts:
            continue

        # Cabeçalho único por via
        header = []
        instit_block = [
            "GOVERNO DO ESTADO DE SÃO PAULO",
            "SECRETARIA DA JUSTIÇA E CIDADANIA",
            "FUNDAÇÃO INSTITUTO DE TERRAS DO ESTADO DE SÃO PAULO",
            '"JOSÉ GOMES DA SILVA" - ITESP',
            "",
            "MEMORIAL DESCRITIVO",
            "",
        ]
        for line in instit_block:
            header.append(f"{line.center(80)}\n")
        header.append(f"VIA: {via_nome}\n")
        header.append(f"BAIRRO: {BAIRRO_FIXO}\n")
        header.append(f"MUNICÍPIO: {MUNICIPIO_FIXO}\n\n")

        descricao_parts = ["DESCRIÇÃO\n\n"]

        for part_idx, segmentos_via in enumerate(via_parts, start=1):
            # Evita iniciar a descrição em ponto de interseção entre vias
            idx_start = next((i for i, s in enumerate(segmentos_via) if "intersectando" not in s["conf"].lower()), None)
            if idx_start is not None and idx_start > 0:
                segmentos_via = segmentos_via[idx_start:] + segmentos_via[:idx_start]

            prefix = f"Segmento {part_idx} - " if len(via_parts) > 1 else ""
            seg_text = []
            seg0 = segmentos_via[0]
            seg_text.append(
                f"{prefix}Inicia-se a descrição desta via no vértice {seg0['from_code']}, "
                f"de coordenadas N {fmt(seg0['from_y'],3)} m e E {fmt(seg0['from_x'],3)} m; "
            )
            for seg in segmentos_via[:-1]:
                conf = seg["conf"]
                if "intersectando" in conf.lower():
                    seg_text.append(
                        f"deste, segue {conf} pela distância de {fmt(seg['dist'],2)} m até o vértice {seg['to_code']}, "
                        f"de coordenadas N {fmt(seg['to_y'],3)} m e E {fmt(seg['to_x'],3)} m; "
                    )
                else:
                    seg_text.append(
                        f"deste, segue {conf} com os seguintes azimutes e distâncias: "
                        f"{seg['az']} e {fmt(seg['dist'],2)} m até o vértice {seg['to_code']}, "
                        f"de coordenadas N {fmt(seg['to_y'],3)} m e E {fmt(seg['to_x'],3)} m; "
                    )
            seg_last = segmentos_via[-1]
            conf_last = seg_last["conf"]
            if "intersectando" in conf_last.lower():
                seg_text.append(
                    f"por fim, segue {conf_last} pela distância de {fmt(seg_last['dist'],2)} m até o vértice {seg_last['to_code']}, "
                    "ponto final desta via. "
                )
            else:
                seg_text.append(
                    f"por fim, segue {conf_last} com os seguintes azimutes e distâncias: "
                    f"{seg_last['az']} e {fmt(seg_last['dist'],2)} m até o vértice {seg_last['to_code']}, "
                    "ponto final desta via. "
                )
            descricao_parts.append("".join(seg_text))
            descricao_parts.append("\n\n")

        descricao_parts.append(
            "Todas as coordenadas aqui descritas estão georreferenciadas ao Sistema "
            "Geodésico Brasileiro, tendo como datum SIRGAS 2000 e sistema de projeção UTM, fuso 22S.\n"
        )

        texto_final = "".join(header) + "".join(descricao_parts)

        if output_fmt == "docx":
            if via_count > 0:
                doc_vias.add_page_break()
            header_lines = block_to_lines(header)
            after_memorial = False
            for line in header_lines:
                p = doc_vias.add_paragraph()
                run = p.add_run(line)
                if line.upper().startswith("MEMORIAL DESCRITIVO"):
                    run.bold = True
                if after_memorial:
                    p.alignment = 0
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if line.upper().startswith("MEMORIAL DESCRITIVO"):
                        after_memorial = True
            for paragraph in "".join(descricao_parts).split("\n\n"):
                text = paragraph.replace("\n", " ").strip()
                if text:
                    p = doc_vias.add_paragraph()
                    run = p.add_run(text)
                    if text.upper().startswith("DESCRIÇÃO"):
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            vias_txt_parts.append(texto_final)
        via_count += 1
        print(f"Via processada: {via_nome}")

    if via_count > 0:
        if output_fmt == "docx":
            doc_vias.save(OUT_VIAS_FILE)
            final_vias_path = OUT_VIAS_FILE
        else:
            final_vias_path = OUT_VIAS_FILE.replace(".docx", ".txt")
            with open(final_vias_path, "w", encoding="utf-8") as f:
                f.write("\n\n".join(vias_txt_parts))
        print(f"Memoriais das vias gerados em: {final_vias_path}")
if segment_records:
    gdf_seg = gpd.GeoDataFrame(segment_records, crs=f"EPSG:{EPSG_CRS}")
    gdf_seg.to_file(OUT_SEGMENTOS_SHP)
    print(f"Shapefile de segmentos (debug) gerado: {OUT_SEGMENTOS_SHP}")
if segment_mid_records:
    gdf_seg_mid = gpd.GeoDataFrame(segment_mid_records, crs=f"EPSG:{EPSG_CRS}")
    gdf_seg_mid.to_file(OUT_SEGMENTOS_PTO_SHP)
    print(f"Shapefile de pontos médios (debug) gerado: {OUT_SEGMENTOS_PTO_SHP}")

print("Processo concluído.")
